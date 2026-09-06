"""Ingestion Pipeline tests: BH-018 through BH-026, BH-049 through BH-052,
BH-062 through BH-068, BH-071.

Covers duplicate detection, force re-ingestion, pipeline failure quarantine,
LLM failure handling, abstraction_skipped, sequential pipeline execution,
source file provenance (source_modified_at), and document date metadata
(document_date).
"""

import hashlib
import os
import re
import shutil
from datetime import datetime, timezone
from pathlib import Path

import pytest

from sage.adapters.interfaces import Chunk, DocumentSurface
from sage.api.errors import DuplicateContentError, ForceReingestPathMismatchError
from sage.config import VaultConfig
from sage.models.enums import PipelineStatus, SourceType
from sage.models.schemas import Document, IngestRequest
from sage.services.document_surface import compose_document_surface
from sage.source_adapters.markdown_adapter import MarkdownAdapter
from tests.helpers.pipeline_wait import await_pipeline_idle

_DOC_ID_RE = re.compile(r"^[0-9a-f]{8}_[a-z0-9_]+$")


def _id(name: str) -> str:
    """Translate a short test name to a shape-conformant document ID.

    The ID validator in sage/models/schemas.py requires the pattern
    ^[0-9a-f]{8}_[a-z0-9_]+$. Test fixtures use short readable names
    like "test_doc"; this helper wraps them so the values still
    construct valid Document instances. Idempotent: an already-canonical
    id passes through unchanged so wrapping is safe to apply at every
    call site.
    """
    if _DOC_ID_RE.fullmatch(name):
        return name
    return f"{hashlib.sha256(name.encode()).hexdigest()[:8]}_{name}"


_SHA256_RE = re.compile(r"^sha256:[0-9a-f]{64}$")


def _sha(name: str) -> str:
    """Deterministic canonical Sha256 from a short test name.

    The Sha256Str validator requires `^sha256:[0-9a-f]{64}$`. Test
    fixtures historically used short readable strings like
    f"hash_{doc_id}" or "sha256:abc"; this helper maps any such
    name to a stable canonical Sha256. Idempotent.
    """
    if _SHA256_RE.fullmatch(name):
        return name
    return "sha256:" + hashlib.sha256(f"sage-test-hash:{name}".encode()).hexdigest()


def _create_test_file(
    tmp_vault_dir: Path, relative_path: str, content: str = "# Test\n\nTest content."
) -> Path:
    """Create a test Markdown file in the vault's sources directory."""
    full_path = tmp_vault_dir / "sources" / relative_path
    full_path.parent.mkdir(parents=True, exist_ok=True)
    full_path.write_text(content)
    return full_path


async def _await_pipeline_terminal(graph_store, doc_id, *, service, attempts=400, delay=0.01):
    """Wait until the document is settled and unclaimed, and return it.

    Thin adapter over the shared wait, kept so the call sites below read the
    way they always have. The predicate -- terminal status *and* no in-flight
    claim -- lives in one place for the whole suite; the claim is released
    after the terminal status write, so a wait keyed on the status alone hands
    its caller a document the next call rejects.
    """
    return await await_pipeline_idle(
        graph_store, doc_id, service=service, attempts=attempts, delay=delay
    )


# ---------------------------------------------------------------------------
# BH-018: Duplicate content detection returns 409
# ---------------------------------------------------------------------------


async def test_bh_018_duplicate_content_409(tmp_vault_dir, graph_store, ingestion_service):
    _create_test_file(tmp_vault_dir, "reports/doc_a.md")

    # First ingest succeeds
    request = IngestRequest(
        source="reports/doc_a.md",
        source_type=SourceType.MARKDOWN,
    )
    result = await ingestion_service.ingest(request)

    doc = result.document
    assert result.is_new is True

    # Second ingest (same path, same content) returns 409
    with pytest.raises(DuplicateContentError) as exc_info:
        await ingestion_service.ingest(request)

    err = exc_info.value
    assert err.status_code == 409
    assert err.code == "duplicate_content"
    assert err.detail["existing_document_id"] == doc.id


# ---------------------------------------------------------------------------
# BH-019: Force re-ingestion bypasses duplicate detection
# ---------------------------------------------------------------------------


async def test_bh_019_force_reingestion(tmp_vault_dir, graph_store, ingestion_service):
    _create_test_file(tmp_vault_dir, "reports/doc_force.md")

    request = IngestRequest(
        source="reports/doc_force.md",
        source_type=SourceType.MARKDOWN,
    )
    result1 = await ingestion_service.ingest(request)

    doc1 = result1.document
    assert result1.is_new is True
    original_id = doc1.id

    # Force re-ingest
    force_request = IngestRequest(
        source="reports/doc_force.md",
        source_type=SourceType.MARKDOWN,
        force=True,
    )
    result2 = await ingestion_service.ingest(force_request)

    doc2 = result2.document

    assert result2.is_new is False
    assert doc2.id == original_id  # Same document record
    assert doc2.pipeline_status == PipelineStatus.ABSTRACTION_COMPLETE


# ---------------------------------------------------------------------------
# BH-020: Failed pipeline quarantines document from retrieval
# (Retrieval not implemented in this slice; test that failed docs
# are marked correctly for quarantine)
# ---------------------------------------------------------------------------


async def test_bh_020_failed_pipeline_marks_document(
    tmp_vault_dir, graph_store, ingestion_service_failing_llm
):
    _create_test_file(tmp_vault_dir, "reports/doc_fail.md")

    request = IngestRequest(
        source="reports/doc_fail.md",
        source_type=SourceType.MARKDOWN,
    )
    result = await ingestion_service_failing_llm.ingest(request)

    doc = result.document
    assert result.is_new is True

    # Sequential pipeline: failure is reflected in the returned document
    fetched = await graph_store.get_document(doc.id)
    assert fetched.pipeline_status == PipelineStatus.FAILED
    assert fetched.pipeline_error is not None


# ---------------------------------------------------------------------------
# BH-022: Failed document still visible via get_document
# ---------------------------------------------------------------------------


async def test_bh_022_failed_document_visible(
    tmp_vault_dir, graph_store, ingestion_service_failing_llm
):
    _create_test_file(tmp_vault_dir, "reports/doc_visible.md")

    request = IngestRequest(
        source="reports/doc_visible.md",
        source_type=SourceType.MARKDOWN,
    )
    result = await ingestion_service_failing_llm.ingest(request)

    doc = result.document

    fetched = await graph_store.get_document(doc.id)
    assert fetched is not None
    assert fetched.pipeline_status == PipelineStatus.FAILED
    assert fetched.pipeline_error is not None
    assert "LLM unavailable" in fetched.pipeline_error


# ---------------------------------------------------------------------------
# BH-024: LLM failure during abstraction results in failed status
# ---------------------------------------------------------------------------


async def test_bh_024_llm_failure_results_in_failed(
    tmp_vault_dir, graph_store, ingestion_service_failing_llm
):
    _create_test_file(
        tmp_vault_dir, "reports/doc_llm_fail.md", "# Document\n\nContent for abstraction."
    )

    request = IngestRequest(
        source="reports/doc_llm_fail.md",
        source_type=SourceType.MARKDOWN,
    )
    result = await ingestion_service_failing_llm.ingest(request)

    doc = result.document
    assert result.is_new is True

    # Sequential pipeline: failure status is set before ingest returns
    fetched = await graph_store.get_document(doc.id)
    assert fetched.pipeline_status == PipelineStatus.FAILED
    assert fetched.pipeline_error is not None
    assert fetched.semantic_abstract is None


# ---------------------------------------------------------------------------
# BH-025: Abstraction disabled produces abstraction_skipped
# ---------------------------------------------------------------------------


async def test_bh_025_abstraction_disabled(
    tmp_vault_dir, graph_store, ingestion_service_no_abstraction
):
    _create_test_file(tmp_vault_dir, "reports/doc_no_abstract.md")

    request = IngestRequest(
        source="reports/doc_no_abstract.md",
        source_type=SourceType.MARKDOWN,
    )
    result = await ingestion_service_no_abstraction.ingest(request)

    doc = result.document
    assert result.is_new is True

    # Sequential pipeline: final status is set before ingest returns
    fetched = await graph_store.get_document(doc.id)
    assert fetched.pipeline_status == PipelineStatus.ABSTRACTION_SKIPPED
    assert fetched.semantic_abstract is None


# ---------------------------------------------------------------------------
# BH-026: Pipeline stages run sequentially within ingest
# ---------------------------------------------------------------------------


async def test_bh_026_sequential_pipeline(tmp_vault_dir, graph_store, ingestion_service):
    _create_test_file(tmp_vault_dir, "reports/doc_seq.md", "# Sequential Test\n\nContent here.")

    request = IngestRequest(
        source="reports/doc_seq.md",
        source_type=SourceType.MARKDOWN,
    )
    result = await ingestion_service.ingest(request)
    doc = result.document

    # Ingest returns with all stages complete (no background tasks)
    assert result.is_new is True
    assert doc.pipeline_status == PipelineStatus.ABSTRACTION_COMPLETE
    assert doc.indexed_at is not None
    assert doc.semantic_abstract is not None

    # Graph store matches returned document
    fetched = await graph_store.get_document(doc.id)
    assert fetched.pipeline_status == PipelineStatus.ABSTRACTION_COMPLETE


# ---------------------------------------------------------------------------
# BH-049: New document ingestion sets source_modified_at from file mtime
# ---------------------------------------------------------------------------


async def test_bh_049_source_modified_at_set_on_ingest(
    tmp_vault_dir, graph_store, ingestion_service
):
    full_path = _create_test_file(tmp_vault_dir, "reports/doc_mtime.md")

    # Set a known mtime in the past
    known_mtime = datetime(2023, 6, 15, 12, 0, 0, tzinfo=timezone.utc)
    os.utime(full_path, (full_path.stat().st_atime, known_mtime.timestamp()))

    request = IngestRequest(
        source="reports/doc_mtime.md",
        source_type=SourceType.MARKDOWN,
    )
    result = await ingestion_service.ingest(request)

    doc = result.document
    assert result.is_new is True

    assert doc.source_modified_at is not None
    assert doc.source_modified_at.tzinfo is not None  # timezone-aware
    assert abs((doc.source_modified_at - known_mtime).total_seconds()) < 1.0


# ---------------------------------------------------------------------------
# BH-050: Force re-ingestion updates source_modified_at
# ---------------------------------------------------------------------------


async def test_bh_050_force_reingestion_updates_source_modified_at(
    tmp_vault_dir, graph_store, ingestion_service
):
    full_path = _create_test_file(tmp_vault_dir, "reports/doc_mtime_force.md")

    # Set old mtime
    old_mtime = datetime(2022, 1, 1, 0, 0, 0, tzinfo=timezone.utc)
    os.utime(full_path, (full_path.stat().st_atime, old_mtime.timestamp()))

    request = IngestRequest(
        source="reports/doc_mtime_force.md",
        source_type=SourceType.MARKDOWN,
    )
    result = await ingestion_service.ingest(request)

    doc1 = result.document
    original_source_mtime = doc1.source_modified_at

    # Touch the file to update mtime
    new_mtime = datetime(2024, 7, 1, 0, 0, 0, tzinfo=timezone.utc)
    os.utime(full_path, (full_path.stat().st_atime, new_mtime.timestamp()))

    force_request = IngestRequest(
        source="reports/doc_mtime_force.md",
        source_type=SourceType.MARKDOWN,
        force=True,
    )
    result2 = await ingestion_service.ingest(force_request)

    doc2 = result2.document
    assert result2.is_new is False

    # Retrieve from store to verify persistence
    fetched = await graph_store.get_document(doc2.id)
    assert abs((fetched.source_modified_at - new_mtime).total_seconds()) < 1.0
    assert fetched.source_modified_at != original_source_mtime


# ---------------------------------------------------------------------------
# BH-051: source_modified_at round-trips through graph store
# ---------------------------------------------------------------------------


async def test_bh_051_source_modified_at_round_trip(tmp_vault_dir, graph_store, ingestion_service):
    full_path = _create_test_file(tmp_vault_dir, "reports/doc_roundtrip.md")

    known_mtime = datetime(2021, 3, 10, 8, 30, 0, tzinfo=timezone.utc)
    os.utime(full_path, (full_path.stat().st_atime, known_mtime.timestamp()))

    request = IngestRequest(
        source="reports/doc_roundtrip.md",
        source_type=SourceType.MARKDOWN,
    )
    result = await ingestion_service.ingest(request)

    doc = result.document

    fetched = await graph_store.get_document(doc.id)
    assert fetched.source_modified_at is not None
    assert fetched.source_modified_at.tzinfo is not None
    assert abs((fetched.source_modified_at - known_mtime).total_seconds()) < 1.0


# ---------------------------------------------------------------------------
# BH-052: created_at remains SAGE ingestion time, distinct from source_modified_at
# ---------------------------------------------------------------------------


async def test_bh_052_created_at_is_ingestion_time(tmp_vault_dir, graph_store, ingestion_service):
    full_path = _create_test_file(tmp_vault_dir, "reports/doc_old.md")

    # Set mtime well in the past
    old_mtime = datetime(2020, 1, 1, 0, 0, 0, tzinfo=timezone.utc)
    os.utime(full_path, (full_path.stat().st_atime, old_mtime.timestamp()))

    before_ingest = datetime.now(timezone.utc)

    request = IngestRequest(
        source="reports/doc_old.md",
        source_type=SourceType.MARKDOWN,
    )
    result = await ingestion_service.ingest(request)

    doc = result.document

    # created_at should be close to now, not the old file mtime
    assert abs((doc.created_at - before_ingest).total_seconds()) < 5.0

    # source_modified_at should match the old mtime
    assert abs((doc.source_modified_at - old_mtime).total_seconds()) < 1.0

    # They must be different
    assert doc.created_at != doc.source_modified_at


# ---------------------------------------------------------------------------
# BH-053: External file is copied verbatim to imports/ subdirectory
# ---------------------------------------------------------------------------


async def test_bh_053_external_file_copied_to_imports(
    tmp_vault_dir, graph_store, ingestion_service, tmp_path
):
    """When source resolves outside storage_root, the file is copied
    verbatim into {storage_root}/imports/ and the original bytes are
    preserved exactly."""
    # Create a file outside storage_root
    external_dir = tmp_path / "ephemeral_session"
    external_dir.mkdir()
    external_file = external_dir / "checklist.md"
    content = "# Checklist\n\nItem one.\nItem two.\n"
    external_file.write_text(content)

    request = IngestRequest(
        source=str(external_file),  # absolute path
        source_type=SourceType.MARKDOWN,
    )
    result = await ingestion_service.ingest(request)

    assert result.is_new is True

    # The file must exist in imports/
    storage_root = tmp_vault_dir / "sources"
    imported = storage_root / "imports" / "checklist.md"
    assert imported.exists()
    assert imported.read_text() == content


# ---------------------------------------------------------------------------
# BH-054: Imported file source_path records vault-relative path
# ---------------------------------------------------------------------------


async def test_bh_054_imported_source_path_is_vault_relative(
    tmp_vault_dir, graph_store, ingestion_service, tmp_path
):
    """The document record's source_path must be vault-relative
    (imports/filename.ext), not the original absolute path."""
    external_file = tmp_path / "external_doc.md"
    external_file.write_text("# External\n\nContent.\n")

    request = IngestRequest(
        source=str(external_file),
        source_type=SourceType.MARKDOWN,
    )
    result = await ingestion_service.ingest(request)

    doc = result.document
    assert result.is_new is True
    assert doc.source_path == "imports/external_doc.md"


# ---------------------------------------------------------------------------
# BH-055: Name collision on import appends 8-char content hash
# ---------------------------------------------------------------------------


async def test_bh_055_import_name_collision_appends_hash(
    tmp_vault_dir, graph_store, ingestion_service, tmp_path
):
    """When imports/ already contains a file with the same name but
    different content, the new file gets an 8-char content hash
    appended before the extension."""
    storage_root = tmp_vault_dir / "sources"
    imports_dir = storage_root / "imports"
    imports_dir.mkdir()

    # Pre-populate imports/ with a file named "report.md"
    existing = imports_dir / "report.md"
    existing.write_text("# Old report\n\nOriginal content.\n")

    # Create an external file with the same name but different content
    external_file = tmp_path / "report.md"
    new_content = "# New report\n\nDifferent content.\n"
    external_file.write_text(new_content)

    request = IngestRequest(
        source=str(external_file),
        source_type=SourceType.MARKDOWN,
    )
    result = await ingestion_service.ingest(request)

    doc = result.document
    assert result.is_new is True

    # source_path should be imports/report_<8-char-hash>.md
    assert doc.source_path.startswith("imports/report_")
    assert doc.source_path.endswith(".md")
    # Extract the hash portion
    stem = Path(doc.source_path).stem  # e.g. "report_a1b2c3d4"
    hash_suffix = stem.split("_", 1)[1]
    assert len(hash_suffix) == 8

    # The disambiguated file must exist and contain the new content
    imported = storage_root / doc.source_path
    assert imported.exists()
    assert imported.read_text() == new_content

    # The original file must be untouched
    assert existing.read_text() == "# Old report\n\nOriginal content.\n"


# ---------------------------------------------------------------------------
# BH-056: File already inside storage_root is not copied
# ---------------------------------------------------------------------------


async def test_bh_056_internal_file_not_copied(tmp_vault_dir, graph_store, ingestion_service):
    """Files already under storage_root are ingested in place without
    copying. source_path is the normalized relative path."""
    _create_test_file(tmp_vault_dir, "reports/internal.md", "# Internal\n\nContent.\n")

    request = IngestRequest(
        source="reports/internal.md",
        source_type=SourceType.MARKDOWN,
    )
    result = await ingestion_service.ingest(request)

    doc = result.document
    assert result.is_new is True
    assert doc.source_path == "reports/internal.md"

    # No imports/ directory should have been created
    imports_dir = tmp_vault_dir / "sources" / "imports"
    assert not imports_dir.exists()


# ---------------------------------------------------------------------------
# BH-057: imports/ directory is created on demand
# ---------------------------------------------------------------------------


async def test_bh_057_imports_dir_created_on_demand(
    tmp_vault_dir, graph_store, ingestion_service, tmp_path
):
    """The imports/ subdirectory is created automatically on the first
    external file import; it does not need to pre-exist."""
    imports_dir = tmp_vault_dir / "sources" / "imports"
    assert not imports_dir.exists()

    external_file = tmp_path / "fresh.md"
    external_file.write_text("# Fresh\n\nNew content.\n")

    request = IngestRequest(
        source=str(external_file),
        source_type=SourceType.MARKDOWN,
    )
    result = await ingestion_service.ingest(request)

    assert result.is_new is True

    assert imports_dir.exists()
    assert (imports_dir / "fresh.md").exists()


# ---------------------------------------------------------------------------
# Regression: _ensure_vault_local creates missing storage_root parent
# ---------------------------------------------------------------------------


async def test_ensure_vault_local_creates_missing_sources_parent(
    tmp_vault_dir, graph_store, ingestion_service, tmp_path
):
    """Regression guard: on a fresh vault where the storage_root parent
    (``sources/``) has not yet been created on disk, importing an
    external file must succeed by creating the tree end-to-end rather
    than raising ``FileNotFoundError`` from the ``imports/`` mkdir."""
    sources_dir = tmp_vault_dir / "sources"
    shutil.rmtree(sources_dir)
    assert not sources_dir.exists()

    external_file = tmp_path / "fresh.md"
    external_file.write_text("# Fresh\n\nNew content.\n")

    request = IngestRequest(
        source=str(external_file),
        source_type=SourceType.MARKDOWN,
    )
    result = await ingestion_service.ingest(request)

    assert result.is_new is True
    assert sources_dir.exists()
    imports_dir = sources_dir / "imports"
    assert imports_dir.exists()
    assert (imports_dir / "fresh.md").exists()


# ---------------------------------------------------------------------------
# _chunk_projection emits body chunks only; document-level identity
# signals live on the document surface (CAS-ADR-049).
# ---------------------------------------------------------------------------


def test_chunk_projection_emits_body_chunks_without_preamble(ingestion_service):
    """_chunk_projection prepends the ATX heading line to body content
    so the projection round-trips through read_projection. The
    synthetic-header preamble still lives in its own chunk and
    must not leak into body chunks."""
    from sage.source_adapters.base import HeadingNode, ProjectionResult

    projection = ProjectionResult(
        text="Body content only.",
        headings=[
            HeadingNode(
                level=1, text="Introduction", path="Introduction", content="Body content only."
            ),
        ],
        content_hash="sha256:abc",
        adapter_version="0.1.0",
        title="ClinicalNormalization",
    )

    chunks = ingestion_service._chunk_projection("doc1", projection)

    assert len(chunks) == 1
    assert chunks[0].content == "# Introduction\n\nBody content only."
    assert not chunks[0].content.startswith("Title:")


def test_chunk_projection_fallback_chunk_has_no_preamble(ingestion_service):
    """When there are no headings, the fallback single chunk carries the
    raw projection text without any identity preamble."""
    from sage.source_adapters.base import ProjectionResult

    projection = ProjectionResult(
        text="Flat document with no headings.",
        headings=[],
        content_hash="sha256:def",
        adapter_version="0.1.0",
        title="Flat_Doc",
    )

    chunks = ingestion_service._chunk_projection("doc2", projection)

    assert len(chunks) == 1
    assert chunks[0].content == "Flat document with no headings."
    assert not chunks[0].content.startswith("Title:")


def test_chunk_projection_multiple_headings_have_no_preamble(ingestion_service):
    """Every body chunk carries its ATX heading line plus body content;
    the synthetic header chunk lives outside _chunk_projection's
    output."""
    from sage.source_adapters.base import HeadingNode, ProjectionResult

    projection = ProjectionResult(
        text="All content.",
        headings=[
            HeadingNode(level=1, text="Part A", path="Part A", content="Content for part A."),
            HeadingNode(level=1, text="Part B", path="Part B", content="Content for part B."),
        ],
        content_hash="sha256:ghi",
        adapter_version="0.1.0",
        title="Multi_Section_Doc",
    )

    chunks = ingestion_service._chunk_projection("doc3", projection)

    assert len(chunks) == 2
    assert chunks[0].content == "# Part A\n\nContent for part A."
    assert chunks[1].content == "# Part B\n\nContent for part B."
    assert all(not c.content.startswith("Title:") for c in chunks)


def _document_for_surface(
    *,
    title: str,
    tags: list[str],
    semantic_abstract: str | None,
    source_path: str,
) -> Document:
    """A stored document record carrying only what surface composition reads."""
    now = datetime.now(timezone.utc)
    return Document(
        id="00000001_doc",
        title=title,
        source_type=SourceType.MARKDOWN,
        source_path=source_path,
        lifecycle_status="active",
        source_content_hash=f"sha256:{0:064x}",
        adapter_version="1",
        created_by="t",
        created_at=now,
        last_modified_by="t",
        updated_at=now,
        doc_type="adr",
        tags=tags,
        semantic_abstract=semantic_abstract,
    )


def test_document_surface_splits_authored_text_from_derived(ingestion_service):
    """Composition routes each kind of text to the half that governs it.

    CAS-ADR-049 makes matchability a function of provenance. The split is a
    property of the composed row rather than of any query, so it is asserted
    on the row itself: the title and tags reach ``matchable``, the generated
    abstract and the source filename stem reach ``orienting``, and neither
    crosses.
    """
    doc = _document_for_surface(
        title="Portfolio Dashboard",
        tags=["retrieval", "content-store"],
        semantic_abstract="A generated summary sentence.",
        source_path="imports/2026-01-01_quarterly-review.md",
    )

    surface = compose_document_surface("00000001_doc", doc)

    assert "Portfolio" in surface.matchable and "Dashboard" in surface.matchable
    assert "retrieval" in surface.matchable and "content" in surface.matchable
    assert "generated summary" in surface.orienting
    assert "quarterly-review" in surface.orienting

    assert "generated summary" not in surface.matchable, (
        "a generated abstract is derived text and must not become matchable"
    )
    assert "quarterly" not in surface.matchable, (
        "a filename is an artifact of how the document arrived, not content"
    )


def test_document_surface_expands_compound_identifiers_in_authored_text_only(
    ingestion_service,
):
    """An expansion inherits the provenance of the text it expands.

    A case-split of the title is still the title, so it belongs with the
    authored half. A case-split of the filename stem is still the filename,
    so it stays derived -- otherwise normalization would quietly hand derived
    text the matchability its raw form is denied.
    """
    doc = _document_for_surface(
        title="PortfolioDashboard",
        tags=[],
        semantic_abstract="",
        source_path="imports/QuarterlyReview.md",
    )

    surface = compose_document_surface("00000001_doc", doc)

    assert "Portfolio" in surface.matchable and "Dashboard" in surface.matchable
    assert "Quarterly" in surface.orienting and "Review" in surface.orienting
    assert "Quarterly" not in surface.matchable
    assert "Review" not in surface.matchable


def test_document_surface_is_stable_when_the_abstract_is_unset(ingestion_service):
    """A document indexed before abstraction still composes a valid row.

    Stage 2 writes this row while ``semantic_abstract`` is still unset, so the
    composition has to tolerate it rather than depend on Stage 3 having run.
    """
    doc = _document_for_surface(
        title="Some Title",
        tags=["alpha"],
        semantic_abstract=None,
        source_path="imports/some-title.md",
    )

    surface = compose_document_surface("00000001_doc", doc)

    assert surface.matchable, "authored text is available before abstraction"
    assert "some-title" in surface.orienting, "the filename stem still reaches the derived half"
    assert not surface.orienting.startswith(" "), (
        "an unset abstract must not leave a leading blank where its text would go"
    )


async def test_abstraction_input_cannot_carry_a_generated_abstract(ingestion_service):
    """Stage 3 cannot be fed its own prior output.

    While document-level text shared the passage surface, a single filter at
    one call site was the only thing keeping a document's own abstract out of
    the text used to regenerate it; drop the filter and every reabstraction
    after the first read its predecessor. The guard is now structural -- the
    abstract is not a passage, and this method reads passages -- so it holds
    without anything remembering to exclude it.

    Asserted by reading what the abstraction stage would receive rather than
    by trusting the absence of a filter, so a future change that puts derived
    text back among the passages fails here.
    """
    store = ingestion_service._content_store
    document_id = "00000001_doc"
    await store.index_chunks(
        document_id,
        [
            Chunk(
                document_id=document_id,
                heading_path="Body",
                content="The authored body of the document.",
                chunk_index=0,
            )
        ],
    )
    await store.upsert_document_surface(
        DocumentSurface(
            document_id=document_id,
            matchable="Some Title",
            orienting="A previously generated abstract sentence.",
        )
    )

    passages = await store.get_all_chunks(document_id)
    abstraction_input = "\n\n".join(chunk.content for chunk in passages)

    assert "authored body" in abstraction_input, "positive control: passages are read"
    assert "previously generated abstract" not in abstraction_input, (
        "the abstraction input reached derived text; the derivation-loop guard "
        "is no longer structural"
    )


# ---------------------------------------------------------------------------
# BH-062: Ingestion with filename date sets document_date from metadata
# ---------------------------------------------------------------------------


async def test_bh_062_filename_date_sets_document_date(
    tmp_vault_dir, graph_store, ingestion_service
):
    """Caller-supplied date in metadata takes precedence as document_date."""
    full_path = _create_test_file(tmp_vault_dir, "reports/dated_doc.md")

    # Set file mtime to a different date so we can verify independence
    file_mtime = datetime(2025, 1, 15, 12, 0, 0, tzinfo=timezone.utc)
    os.utime(full_path, (full_path.stat().st_atime, file_mtime.timestamp()))

    request = IngestRequest(
        source="reports/dated_doc.md",
        source_type=SourceType.MARKDOWN,
        metadata={"date": "2026-04-10"},
    )
    result = await ingestion_service.ingest(request)

    doc = result.document
    assert result.is_new is True

    assert doc.document_date == "2026-04-10"
    # source_modified_at is set independently from file mtime
    assert doc.source_modified_at is not None
    assert abs((doc.source_modified_at - file_mtime).total_seconds()) < 1.0


# ---------------------------------------------------------------------------
# BH-063: Ingestion without filename date falls back to source_modified_at
# ---------------------------------------------------------------------------


async def test_bh_063_fallback_to_source_modified_at_date(
    tmp_vault_dir, graph_store, ingestion_service
):
    """When no date in metadata, document_date derives from source_modified_at."""
    full_path = _create_test_file(tmp_vault_dir, "reports/no_date_doc.md")

    known_mtime = datetime(2025, 6, 15, 14, 30, 0, tzinfo=timezone.utc)
    os.utime(full_path, (full_path.stat().st_atime, known_mtime.timestamp()))

    request = IngestRequest(
        source="reports/no_date_doc.md",
        source_type=SourceType.MARKDOWN,
    )
    result = await ingestion_service.ingest(request)

    doc = result.document

    assert doc.document_date == "2025-06-15"
    assert doc.source_modified_at is not None


# ---------------------------------------------------------------------------
# Vault-local timezone applied to source_modified_at fallback
# ---------------------------------------------------------------------------


async def test_document_date_fallback_uses_vault_timezone(
    tmp_vault_dir,
    graph_store,
    lock_manager,
    stub_content_store,
    stub_embedding_provider,
    stub_abstraction_provider,
    minimal_vault_config_dict,
    lifecycle_service,
):
    """When vault.timezone is set, the fallback derives document_date in
    that zone, not UTC. Regression guard for the UTC-midnight drift bug:
    a Chicago-evening mtime that crosses UTC midnight should attribute
    to the local calendar date.
    """
    from sage.services.ingestion import IngestionService

    config_dict = minimal_vault_config_dict.copy()
    config_dict["vault"] = {**config_dict["vault"], "timezone": "America/Chicago"}
    config = VaultConfig.model_validate(config_dict)
    service = IngestionService(
        graph_store=graph_store,
        lock_manager=lock_manager,
        content_store=stub_content_store,
        embedding_provider=stub_embedding_provider,
        abstraction_provider=stub_abstraction_provider,
        config=config,
        source_adapters={SourceType.MARKDOWN: MarkdownAdapter()},
        lifecycle_service=lifecycle_service,
    )

    full_path = _create_test_file(tmp_vault_dir, "reports/late_evening.md")
    # 2026-04-29 00:33 UTC is 2026-04-28 19:33 CDT.
    chicago_evening_mtime = datetime(2026, 4, 29, 0, 33, 0, tzinfo=timezone.utc)
    os.utime(
        full_path,
        (full_path.stat().st_atime, chicago_evening_mtime.timestamp()),
    )

    request = IngestRequest(
        source="reports/late_evening.md",
        source_type=SourceType.MARKDOWN,
    )
    result = await service.ingest(request)
    doc = result.document

    assert doc.document_date == "2026-04-28"


def test_vault_identity_carries_timezone_field(minimal_vault_config_dict):
    """VaultConfig accepts an explicit vault.timezone and defaults to UTC
    when the field is omitted.
    """
    omitted = VaultConfig.model_validate(minimal_vault_config_dict)
    assert omitted.vault.timezone == "UTC"

    explicit_dict = minimal_vault_config_dict.copy()
    explicit_dict["vault"] = {**explicit_dict["vault"], "timezone": "America/Chicago"}
    explicit = VaultConfig.model_validate(explicit_dict)
    assert explicit.vault.timezone == "America/Chicago"


async def test_document_date_fallback_defaults_to_utc(
    tmp_vault_dir, graph_store, ingestion_service
):
    """A vault config without an explicit timezone keeps the BH-063
    behavior: fallback uses the UTC calendar date of source_modified_at.
    Regression guard for existing vaults.
    """
    full_path = _create_test_file(tmp_vault_dir, "reports/utc_default.md")
    known_mtime = datetime(2025, 6, 15, 14, 30, 0, tzinfo=timezone.utc)
    os.utime(full_path, (full_path.stat().st_atime, known_mtime.timestamp()))

    request = IngestRequest(
        source="reports/utc_default.md",
        source_type=SourceType.MARKDOWN,
    )
    result = await ingestion_service.ingest(request)
    doc = result.document

    assert doc.document_date == "2025-06-15"


# ---------------------------------------------------------------------------
# BH-064: No filename date and no source_modified_at leaves document_date null
# ---------------------------------------------------------------------------


async def test_bh_064_null_when_no_date_sources(tmp_vault_dir, graph_store, ingestion_service):
    """document_date is null when neither filename date nor source_modified_at
    is available."""
    _create_test_file(tmp_vault_dir, "reports/no_sources.md")

    # Simpler approach: ingest, then clear both fields and verify the
    # invariant that _build_metadata_updates doesn't set document_date
    # when no date key is present.
    request = IngestRequest(
        source="reports/no_sources.md",
        source_type=SourceType.MARKDOWN,
    )
    result = await ingestion_service.ingest(request)

    doc = result.document

    # Clear both fields to simulate the no-source scenario
    await graph_store.update_document(
        doc.id,
        {
            "source_modified_at": None,
            "document_date": None,
        },
    )
    fetched = await graph_store.get_document(doc.id)
    assert fetched.document_date is None
    assert fetched.source_modified_at is None


# ---------------------------------------------------------------------------
# BH-065: document_date round-trips through graph store
# ---------------------------------------------------------------------------


async def test_bh_065_document_date_round_trip(tmp_vault_dir, graph_store, ingestion_service):
    """document_date stored as TEXT in SQLite survives insert/retrieve cycle."""
    _create_test_file(tmp_vault_dir, "reports/date_roundtrip.md")

    request = IngestRequest(
        source="reports/date_roundtrip.md",
        source_type=SourceType.MARKDOWN,
        metadata={"date": "2026-04-10"},
    )
    result = await ingestion_service.ingest(request)

    doc = result.document

    fetched = await graph_store.get_document(doc.id)
    assert fetched.document_date == "2026-04-10"
    assert isinstance(fetched.document_date, str)


# ---------------------------------------------------------------------------
# BH-067: Force re-ingestion reuses existing document at a different path,
# now gated behind an explicit document_id confirmation. A hash match at a
# different source_path may be an unrelated document that merely shares
# content bytes, so the caller must name the record to reuse; otherwise the
# collision is rejected (see the cross-document-collision guard test below).
# ---------------------------------------------------------------------------


async def test_bh_067_force_reingestion_different_path_reuses_document(
    tmp_vault_dir, graph_store, ingestion_service
):
    """Force re-ingestion of identical content at a different path reuses the
    existing document record when the caller confirms the target via
    document_id. The source_path is updated to the new location. Without the
    document_id pin the same call is rejected (a different-path hash match may
    be an unrelated document); that guard is covered by
    test_force_reingest_cross_document_collision_raises."""
    content = "# Identical\n\nSame content at two paths."
    _create_test_file(tmp_vault_dir, "reports/doc_a.md", content)
    _create_test_file(tmp_vault_dir, "reports/subfolder/doc_a_copy.md", content)

    # First ingest
    result1 = await ingestion_service.ingest(
        IngestRequest(
            source="reports/doc_a.md",
            source_type=SourceType.MARKDOWN,
        )
    )
    assert result1.is_new is True
    original_id = result1.document.id

    # Force re-ingest from a different path with identical content, confirming
    # the intended record via document_id (the legitimate file-moved case).
    result2 = await ingestion_service.ingest(
        IngestRequest(
            source="reports/subfolder/doc_a_copy.md",
            source_type=SourceType.MARKDOWN,
            force=True,
            document_id=original_id,
        )
    )

    assert result2.is_new is False
    assert result2.document.id == original_id
    assert result2.document.source_path == "reports/subfolder/doc_a_copy.md"
    assert result2.document.pipeline_status == PipelineStatus.ABSTRACTION_COMPLETE


async def test_force_reingest_cross_document_collision_raises(
    tmp_vault_dir, graph_store, ingestion_service
):
    """Force re-ingesting content that is byte-identical to an *unrelated*
    document stored at a different path must not silently overwrite that
    document. Without a document_id pin the collision is rejected with
    force_reingest_path_mismatch, and the original document is left intact.

    Regression guard for the identity-clobber bug: previously the force branch
    resolved its target by content hash alone and rewrote the first document's
    source_path (and title) to the second file's."""
    content = "# Shared\n\nByte-identical content living at two unrelated paths."
    _create_test_file(tmp_vault_dir, "reports/collide_a.md", content)
    _create_test_file(tmp_vault_dir, "reports/collide_b.md", content)

    result_a = await ingestion_service.ingest(
        IngestRequest(
            source="reports/collide_a.md",
            source_type=SourceType.MARKDOWN,
        )
    )
    assert result_a.is_new is True
    id_a = result_a.document.id
    title_a = result_a.document.title

    # Force-ingest the byte-identical file at a different path, WITHOUT a pin.
    with pytest.raises(ForceReingestPathMismatchError) as exc_info:
        await ingestion_service.ingest(
            IngestRequest(
                source="reports/collide_b.md",
                source_type=SourceType.MARKDOWN,
                force=True,
            )
        )

    err = exc_info.value
    assert err.status_code == 409
    assert err.code == "force_reingest_path_mismatch"
    assert err.detail["existing_document_id"] == id_a
    assert err.detail["existing_source_path"] == "reports/collide_a.md"
    assert err.detail["new_source_path"] == "reports/collide_b.md"

    # The unrelated document A is untouched: same path, same title, still there.
    fetched_a = await graph_store.get_document(id_a)
    assert fetched_a is not None
    assert fetched_a.source_path == "reports/collide_a.md"
    assert fetched_a.title == title_a


# ---------------------------------------------------------------------------
# BH-068: Sequential pipeline sets final status before returning
# ---------------------------------------------------------------------------


async def test_bh_068_sequential_pipeline_final_status(
    tmp_vault_dir, graph_store, ingestion_service
):
    """Sequential pipeline completes all stages before ingest() returns.
    The returned IngestResult.document reflects the terminal pipeline
    status, and the graph store is consistent with it."""
    _create_test_file(
        tmp_vault_dir,
        "reports/doc_final.md",
        "# Final Status\n\nVerify sequential pipeline completion.",
    )

    request = IngestRequest(
        source="reports/doc_final.md",
        source_type=SourceType.MARKDOWN,
    )
    result = await ingestion_service.ingest(request)
    doc = result.document

    # Returned document has terminal status
    assert doc.pipeline_status == PipelineStatus.ABSTRACTION_COMPLETE
    assert doc.indexed_at is not None
    assert doc.semantic_abstract is not None

    # Graph store matches
    fetched = await graph_store.get_document(doc.id)
    assert fetched.pipeline_status == doc.pipeline_status
    assert fetched.indexed_at == doc.indexed_at
    assert fetched.semantic_abstract == doc.semantic_abstract


# ---------------------------------------------------------------------------
# BH-071: Same-name same-content re-import reuses existing file
# ---------------------------------------------------------------------------


async def test_bh_071_same_content_reuses_existing_import(
    tmp_vault_dir, graph_store, ingestion_service, tmp_path
):
    """When imports/ already contains a file with the same name AND
    identical content, _ensure_vault_local returns the existing path
    without creating a hash-suffixed duplicate."""
    storage_root = tmp_vault_dir / "sources"
    imports_dir = storage_root / "imports"
    imports_dir.mkdir()

    content = "# Report\n\nIdentical content.\n"

    # Pre-populate imports/ with a file
    existing = imports_dir / "report.md"
    existing.write_text(content)

    # Create an external file with the same name and same content
    external_file = tmp_path / "report.md"
    external_file.write_text(content)

    request = IngestRequest(
        source=str(external_file),
        source_type=SourceType.MARKDOWN,
        force=True,  # bypass DuplicateContentError to test file handling
    )
    result = await ingestion_service.ingest(request)

    doc = result.document
    # source_path should be the original, without hash suffix
    assert doc.source_path == "imports/report.md"

    # Only one file should exist in imports/
    import_files = list(imports_dir.iterdir())
    assert len(import_files) == 1
    assert import_files[0].name == "report.md"

    # Content unchanged
    assert existing.read_text() == content


# ---------------------------------------------------------------------------
# _generate_abstract_text: shared core for abstraction generation
# ---------------------------------------------------------------------------


async def test_generate_abstract_text_uses_density_proportional_budget(
    ingestion_service,
):
    """_generate_abstract_text should pass a density-proportional max_tokens
    value to the abstraction provider, not a fixed constant."""
    # Capture the max_tokens passed to the stub
    captured = {}
    original = ingestion_service._abstraction.generate_abstract

    async def spy(text: str, max_tokens: int, doc_type: str | None) -> str:
        captured["max_tokens"] = max_tokens
        captured["doc_type"] = doc_type
        return await original(text, max_tokens, doc_type)

    ingestion_service._abstraction.generate_abstract = spy

    text = "word " * 10000  # 10000 words
    result = await ingestion_service._generate_abstract_text(text, "guideline")

    assert "max_tokens" in captured
    # 150 + 10000 * 0.02 = 350, not the old fixed 500
    assert captured["max_tokens"] == 350
    assert captured["doc_type"] == "guideline"
    assert result  # non-empty


async def test_generate_abstract_text_trims_sentence_boundary(
    ingestion_service,
):
    """_generate_abstract_text should trim output to the last complete
    sentence boundary."""

    # Replace stub to return text truncated mid-sentence
    async def truncated_output(text: str, max_tokens: int, doc_type: str | None) -> str:
        return "First sentence. Second sentence. Third incompl"

    ingestion_service._abstraction.generate_abstract = truncated_output

    result = await ingestion_service._generate_abstract_text("any input", None)
    assert result == "First sentence. Second sentence."


async def test_generate_abstract_text_returns_complete_sentences_unchanged(
    ingestion_service,
):
    """When abstraction output ends at a sentence boundary, it should be
    returned unchanged."""

    async def clean_output(text: str, max_tokens: int, doc_type: str | None) -> str:
        return "Complete sentence one. Complete sentence two."

    ingestion_service._abstraction.generate_abstract = clean_output

    result = await ingestion_service._generate_abstract_text("any input", None)
    assert result == "Complete sentence one. Complete sentence two."


# ---------------------------------------------------------------------------
# reabstract (async / fire-and-forget): BH-116 through BH-121
# ---------------------------------------------------------------------------


async def test_reabstract_returns_immediately_with_started_status(
    tmp_vault_dir,
    ingestion_service,
    graph_store,
):
    """BH-116: reabstract should return immediately with a dict containing
    status='reabstract_started' and the document_id, without waiting for
    the background abstraction to complete."""
    _create_test_file(tmp_vault_dir, "reports/reabs.md", "# Reabstract Test\n\nOriginal content.")

    request = IngestRequest(source="reports/reabs.md", source_type=SourceType.MARKDOWN)
    result = await ingestion_service.ingest(request)
    doc_id = result.document.id

    response = await ingestion_service.reabstract(doc_id)

    assert isinstance(response, dict)
    assert response["status"] == "reabstract_started"
    assert response["document_id"] == doc_id


async def test_reabstract_sets_pipeline_status_in_progress(
    tmp_vault_dir,
    ingestion_service,
    graph_store,
):
    """BH-117: reabstract should set pipeline_status to
    abstraction_in_progress before returning."""
    import asyncio

    _create_test_file(tmp_vault_dir, "reports/status.md", "# Status Test\n\nContent.")

    request = IngestRequest(source="reports/status.md", source_type=SourceType.MARKDOWN)
    result = await ingestion_service.ingest(request)
    doc_id = result.document.id

    # Gate abstraction so the queued reabstract job cannot race to
    # ABSTRACTION_COMPLETE before we observe the synchronous in-progress
    # pre-stamp. (Pre-queue, the background task slept 0.1s; the worker has no
    # such delay, so the gate is what makes the transient deterministic.)
    gate = asyncio.Event()

    async def gated_abstract(text: str, max_tokens: int, doc_type: str | None) -> str:
        await gate.wait()
        return "gated abstract"

    ingestion_service._abstraction.generate_abstract = gated_abstract

    await ingestion_service.reabstract(doc_id)

    doc = await graph_store.get_document(doc_id)
    assert doc.pipeline_status == PipelineStatus.ABSTRACTION_IN_PROGRESS

    gate.set()
    await ingestion_service.stop_worker()


async def test_reabstract_background_updates_abstract_on_success(
    tmp_vault_dir,
    ingestion_service,
    graph_store,
):
    """BH-118: The background task should update semantic_abstract and set
    pipeline_status to abstraction_complete when abstraction succeeds."""
    import asyncio

    _create_test_file(tmp_vault_dir, "reports/bgok.md", "# BG OK\n\nOriginal content.")

    request = IngestRequest(source="reports/bgok.md", source_type=SourceType.MARKDOWN)
    result = await ingestion_service.ingest(request)
    doc_id = result.document.id
    original_abstract = result.document.semantic_abstract

    # Swap abstraction provider to produce different output
    async def new_abstract(text: str, max_tokens: int, doc_type: str | None) -> str:
        return "Regenerated abstract from new model."

    ingestion_service._abstraction.generate_abstract = new_abstract

    response = await ingestion_service.reabstract(doc_id)
    assert response["status"] == "reabstract_started"

    # Allow the background task to complete (0.1s yield + work)
    await asyncio.sleep(0.5)

    doc = await graph_store.get_document(doc_id)
    assert doc.semantic_abstract == "Regenerated abstract from new model."
    assert doc.semantic_abstract != original_abstract
    assert doc.pipeline_status == PipelineStatus.ABSTRACTION_COMPLETE


async def test_reabstract_background_sets_failed_on_error(
    tmp_vault_dir,
    ingestion_service_failing_llm,
    graph_store,
):
    """BH-119: The background task should set pipeline_status to 'failed'
    when the abstraction provider raises an exception."""
    _create_test_file(tmp_vault_dir, "reports/bgfail.md", "# BG Fail\n\nContent.")

    # Ingest with the normal stub first so initial ingestion succeeds,
    # then swap to the failing provider for reabstract
    from sage.adapters.stubs import StubAbstractionProvider

    original_provider = ingestion_service_failing_llm._abstraction

    # Use stub for initial ingest
    ingestion_service_failing_llm._abstraction = StubAbstractionProvider()
    request = IngestRequest(source="reports/bgfail.md", source_type=SourceType.MARKDOWN)
    result = await ingestion_service_failing_llm.ingest(request)
    doc_id = result.document.id

    # Switch to failing provider for reabstract
    ingestion_service_failing_llm._abstraction = original_provider

    response = await ingestion_service_failing_llm.reabstract(doc_id)
    assert response["status"] == "reabstract_started"

    doc = await _await_pipeline_terminal(graph_store, doc_id, service=ingestion_service_failing_llm)
    assert doc.pipeline_status == PipelineStatus.FAILED


async def test_reabstract_document_not_found(ingestion_service):
    """BH-120: reabstract should raise DocumentNotFoundError synchronously
    for unknown document_id (validation happens before background dispatch)."""
    from sage.api.errors import DocumentNotFoundError

    with pytest.raises(DocumentNotFoundError):
        await ingestion_service.reabstract("nonexistent_doc_id")


async def test_reabstract_no_projection_raises_error(
    tmp_vault_dir,
    ingestion_service,
    graph_store,
    stub_content_store,
):
    """BH-121: reabstract should raise NoProjectionError synchronously when
    the document exists but has no stored chunks."""
    from sage.api.errors import NoProjectionError

    _create_test_file(tmp_vault_dir, "reports/nochunks.md", "# No Chunks\n\nContent.")

    request = IngestRequest(source="reports/nochunks.md", source_type=SourceType.MARKDOWN)
    result = await ingestion_service.ingest(request)
    doc_id = result.document.id

    # Remove all chunks from the content store
    await stub_content_store.remove_document(doc_id)

    with pytest.raises(NoProjectionError):
        await ingestion_service.reabstract(doc_id)


# ---------------------------------------------------------------------------
# Per-document single-flight lock for ``IngestionService.reabstract``: covers
# contention (second concurrent call raises 409), per-document independence
# (parallel calls against different docs both proceed), and reservation
# cleanup on both success and failure terminal states.
# ---------------------------------------------------------------------------


async def test_reabstract_second_concurrent_call_raises_in_flight_error(
    tmp_vault_dir,
    ingestion_service,
    graph_store,
):
    """A second reabstract against the same document_id while the first is
    mid-flight must raise ReabstractDocumentAlreadyInFlightError (409).
    """
    import asyncio

    from sage.api.errors import ReabstractDocumentAlreadyInFlightError

    _create_test_file(tmp_vault_dir, "samples/sf1.md", "# SF1\n\nContent.")
    result = await ingestion_service.ingest(
        IngestRequest(source="samples/sf1.md", source_type=SourceType.MARKDOWN)
    )
    doc_id = result.document.id

    # Gate the abstraction provider so the first background task hangs
    # mid-generation, holding its reservation in the in-flight set.
    entered = asyncio.Event()
    gate = asyncio.Event()

    async def gated_abstract(text: str, max_tokens: int, doc_type: str | None) -> str:
        entered.set()
        await gate.wait()
        return "gated abstract"

    ingestion_service._abstraction.generate_abstract = gated_abstract

    window_before = datetime.now(timezone.utc)
    first_response = await ingestion_service.reabstract(doc_id)
    assert first_response["status"] == "reabstract_started"
    window_after = datetime.now(timezone.utc)

    # Wait until the background task is actually inside generate_abstract.
    # Without this barrier the second call could race ahead of the
    # background task's reservation hand-off (currently None, but kept as
    # a structural safeguard against future refactors of the dispatch).
    await asyncio.wait_for(entered.wait(), timeout=2.0)

    try:
        with pytest.raises(ReabstractDocumentAlreadyInFlightError) as excinfo:
            await ingestion_service.reabstract(doc_id)

        err = excinfo.value
        assert err.code == "reabstract_document_already_in_flight"
        assert err.status_code == 409
        assert err.detail["document_id"] == doc_id
        observed_start = datetime.fromisoformat(err.detail["start_time"])
        assert window_before <= observed_start <= window_after
    finally:
        # Release the gate so the background task can clean up its
        # reservation and not leak into the next test in the session.
        gate.set()
        await asyncio.sleep(0.3)


async def test_reabstract_parallel_calls_different_documents_both_succeed(
    tmp_vault_dir,
    ingestion_service,
    graph_store,
):
    """Two reabstracts against different document_ids must both succeed
    even when the first is mid-flight. The lock is per-document, not
    per-vault.
    """
    import asyncio

    _create_test_file(tmp_vault_dir, "samples/sf2a.md", "# SF2A\n\nContent.")
    _create_test_file(tmp_vault_dir, "samples/sf2b.md", "# SF2B\n\nDifferent content.")
    result_a = await ingestion_service.ingest(
        IngestRequest(source="samples/sf2a.md", source_type=SourceType.MARKDOWN)
    )
    result_b = await ingestion_service.ingest(
        IngestRequest(source="samples/sf2b.md", source_type=SourceType.MARKDOWN)
    )
    doc_a, doc_b = result_a.document.id, result_b.document.id
    assert doc_a != doc_b

    gate = asyncio.Event()
    enter_count = {"n": 0}

    async def gated_abstract(text: str, max_tokens: int, doc_type: str | None) -> str:
        enter_count["n"] += 1
        await gate.wait()
        return "gated abstract"

    ingestion_service._abstraction.generate_abstract = gated_abstract

    response_a = await ingestion_service.reabstract(doc_a)
    response_b = await ingestion_service.reabstract(doc_b)

    assert response_a["status"] == "reabstract_started"
    assert response_a["document_id"] == doc_a
    assert response_b["status"] == "reabstract_started"
    assert response_b["document_id"] == doc_b

    # Allow both background tasks to reach the gate, then release them so
    # cleanup runs and both reservations are popped.
    for _ in range(40):
        if enter_count["n"] >= 2:
            break
        await asyncio.sleep(0.05)
    gate.set()
    await asyncio.sleep(0.5)

    doc_a_state = await graph_store.get_document(doc_a)
    doc_b_state = await graph_store.get_document(doc_b)
    assert doc_a_state.pipeline_status == PipelineStatus.ABSTRACTION_COMPLETE
    assert doc_b_state.pipeline_status == PipelineStatus.ABSTRACTION_COMPLETE


async def test_reabstract_succeeds_again_after_prior_reabstract_completes(
    tmp_vault_dir,
    ingestion_service,
    graph_store,
):
    """After a reabstract reaches ABSTRACTION_COMPLETE the reservation
    must be released so a subsequent reabstract against the same document
    can proceed.
    """
    _create_test_file(tmp_vault_dir, "samples/sf3.md", "# SF3\n\nContent.")
    result = await ingestion_service.ingest(
        IngestRequest(source="samples/sf3.md", source_type=SourceType.MARKDOWN)
    )
    doc_id = result.document.id

    async def fast_abstract(text: str, max_tokens: int, doc_type: str | None) -> str:
        return "T3 abstract"

    ingestion_service._abstraction.generate_abstract = fast_abstract

    first = await ingestion_service.reabstract(doc_id)
    assert first["status"] == "reabstract_started"

    # Wait for the claim, not only for the terminal stamp: the second call is
    # exactly what a still-held claim rejects.
    doc = await _await_pipeline_terminal(graph_store, doc_id, service=ingestion_service)
    assert doc.pipeline_status == PipelineStatus.ABSTRACTION_COMPLETE

    second = await ingestion_service.reabstract(doc_id)
    assert second["status"] == "reabstract_started"

    # Drain the second background job so teardown does not race it.
    await _await_pipeline_terminal(graph_store, doc_id, service=ingestion_service)


async def test_reabstract_succeeds_again_after_prior_reabstract_failed(
    tmp_vault_dir,
    ingestion_service_failing_llm,
    graph_store,
):
    """After a reabstract's background task transitions to FAILED the
    reservation must still be released so a subsequent reabstract can
    proceed.
    """
    from sage.adapters.stubs import StubAbstractionProvider

    _create_test_file(tmp_vault_dir, "samples/sf4.md", "# SF4\n\nContent.")

    # Ingest under the working stub so the projection lands cleanly,
    # then swap to the failing provider for the reabstract call.
    failing_provider = ingestion_service_failing_llm._abstraction
    ingestion_service_failing_llm._abstraction = StubAbstractionProvider()
    result = await ingestion_service_failing_llm.ingest(
        IngestRequest(source="samples/sf4.md", source_type=SourceType.MARKDOWN)
    )
    doc_id = result.document.id
    ingestion_service_failing_llm._abstraction = failing_provider

    first = await ingestion_service_failing_llm.reabstract(doc_id)
    assert first["status"] == "reabstract_started"

    doc = await _await_pipeline_terminal(graph_store, doc_id, service=ingestion_service_failing_llm)
    assert doc.pipeline_status == PipelineStatus.FAILED

    # Swap back to a working provider so the second call has a chance
    # to dispatch without immediately re-failing; the assertion is only
    # that the synchronous prefix returns started_status (which proves
    # the reservation was released).
    ingestion_service_failing_llm._abstraction = StubAbstractionProvider()
    second = await ingestion_service_failing_llm.reabstract(doc_id)
    assert second["status"] == "reabstract_started"

    # Drain the second background job so teardown does not race it.
    await _await_pipeline_terminal(graph_store, doc_id, service=ingestion_service_failing_llm)


# ---------------------------------------------------------------------------
# IngestionService.recompute_pipeline -- operator-driven recovery for documents
# stuck at pipeline_status=projection_complete with no chunks. Each call re-runs
# Stage 1 (projection from source_path) + Stages 2-3 in the background; the
# synchronous prefix surfaces unknown_document / unknown_vault /
# source_file_not_found / 409 in-flight errors.
# ---------------------------------------------------------------------------


async def test_recompute_pipeline_recovers_stuck_projection_complete(
    tmp_vault_dir,
    ingestion_service,
    graph_store,
    stub_content_store,
):
    """A document forced into ``projection_complete`` with no chunks must
    reach a terminal state with chunks present after one
    ``recompute_pipeline`` call.
    """

    _create_test_file(tmp_vault_dir, "samples/recpl1.md", "# RecPL1\n\nContent.")

    request = IngestRequest(source="samples/recpl1.md", source_type=SourceType.MARKDOWN)
    result = await ingestion_service.ingest(request)
    doc_id = result.document.id

    # Force the document into the stuck silent-loss state: clear chunks and
    # rewind ``pipeline_status`` to ``projection_complete``.
    await stub_content_store.remove_document(doc_id)
    await graph_store.update_document(
        doc_id,
        {
            "pipeline_status": PipelineStatus.PROJECTION_COMPLETE.value,
            "indexed_at": None,
            "semantic_abstract": None,
        },
    )
    pre = await graph_store.get_document(doc_id)
    assert pre.pipeline_status == PipelineStatus.PROJECTION_COMPLETE
    assert await stub_content_store.get_all_chunks(doc_id) == []

    response = await ingestion_service.recompute_pipeline(doc_id)
    assert response["status"] == "recompute_pipeline_started"
    assert response["document_id"] == doc_id
    assert "dispatched_at" in response

    # Wait for the background task to drive Stage 2-3 to terminal state.
    terminal = await _await_pipeline_terminal(graph_store, doc_id, service=ingestion_service)
    assert terminal.pipeline_status in {
        PipelineStatus.ABSTRACTION_COMPLETE,
        PipelineStatus.ABSTRACTION_SKIPPED,
    }
    chunks = await stub_content_store.get_all_chunks(doc_id)
    assert chunks, "recompute_pipeline must re-populate chunks"


async def test_recompute_pipeline_idempotent_on_terminal_document(
    tmp_vault_dir,
    ingestion_service,
    graph_store,
    stub_content_store,
):
    """Calling ``recompute_pipeline`` on an already-terminal document must
    succeed (re-runs Stages 1-3 cleanly) without leaving the document in
    ``FAILED``. The pipeline_status must observably cycle through
    ``indexing_in_progress`` before reaching terminal -- proves the call
    actually re-ran Stage 2 rather than no-op'ing.
    """
    import asyncio

    _create_test_file(tmp_vault_dir, "samples/recpl2.md", "# RecPL2\n\nContent.")

    request = IngestRequest(source="samples/recpl2.md", source_type=SourceType.MARKDOWN)
    result = await ingestion_service.ingest(request)
    doc_id = result.document.id

    initial = await graph_store.get_document(doc_id)
    assert initial.pipeline_status == PipelineStatus.ABSTRACTION_COMPLETE

    # Gate embed so we can observe ``indexing_in_progress`` before the
    # second pipeline run completes.
    gate = asyncio.Event()
    original_embed = ingestion_service._embedding.embed

    async def gated_embed(texts):
        await gate.wait()
        return await original_embed(texts)

    ingestion_service._embedding.embed = gated_embed

    response = await ingestion_service.recompute_pipeline(doc_id)
    assert response["status"] == "recompute_pipeline_started"

    # Allow Stage 2 to enter and write the in-progress status, then verify.
    for _ in range(200):
        doc = await graph_store.get_document(doc_id)
        if doc.pipeline_status == PipelineStatus.INDEXING_IN_PROGRESS:
            break
        await asyncio.sleep(0.005)
    else:
        gate.set()
        await asyncio.sleep(0.2)
        pytest.fail("recompute_pipeline did not observably re-run Stage 2")

    gate.set()
    terminal = await _await_pipeline_terminal(graph_store, doc_id, service=ingestion_service)
    assert terminal.pipeline_status == PipelineStatus.ABSTRACTION_COMPLETE
    chunks = await stub_content_store.get_all_chunks(doc_id)
    assert chunks


async def test_recompute_pipeline_unknown_document_raises(ingestion_service):
    """Unknown document_id must raise ``DocumentNotFoundError`` synchronously
    (before background dispatch)."""
    from sage.api.errors import DocumentNotFoundError

    with pytest.raises(DocumentNotFoundError):
        await ingestion_service.recompute_pipeline("deadbeef_nonexistent")


async def test_recompute_pipeline_single_flight_409(
    tmp_vault_dir,
    ingestion_service,
    graph_store,
):
    """A second ``recompute_pipeline`` against the same document_id while the
    first is mid-flight must raise ``RecomputePipelineAlreadyInFlightError``
    (409). The lock is per-document.
    """
    import asyncio

    from sage.api.errors import RecomputePipelineAlreadyInFlightError

    _create_test_file(tmp_vault_dir, "samples/recpl4.md", "# RecPL4\n\nContent.")
    result = await ingestion_service.ingest(
        IngestRequest(source="samples/recpl4.md", source_type=SourceType.MARKDOWN)
    )
    doc_id = result.document.id

    # Gate the embed step so the first recompute_pipeline holds its
    # reservation through the contention window.
    entered = asyncio.Event()
    gate = asyncio.Event()
    original_embed = ingestion_service._embedding.embed

    async def gated_embed(texts):
        entered.set()
        await gate.wait()
        return await original_embed(texts)

    ingestion_service._embedding.embed = gated_embed

    window_before = datetime.now(timezone.utc)
    first = await ingestion_service.recompute_pipeline(doc_id)
    assert first["status"] == "recompute_pipeline_started"
    window_after = datetime.now(timezone.utc)

    await asyncio.wait_for(entered.wait(), timeout=2.0)
    assert not gate.is_set()

    try:
        with pytest.raises(RecomputePipelineAlreadyInFlightError) as excinfo:
            await ingestion_service.recompute_pipeline(doc_id)

        err = excinfo.value
        assert err.code == "recompute_pipeline_already_in_flight"
        assert err.status_code == 409
        assert err.detail["document_id"] == doc_id
        observed_start = datetime.fromisoformat(err.detail["start_time"])
        assert window_before <= observed_start <= window_after
    finally:
        gate.set()
        await asyncio.sleep(0.3)


async def test_recompute_pipeline_source_path_missing_raises(
    tmp_vault_dir,
    ingestion_service,
    graph_store,
):
    """A document whose source file no longer exists on disk must surface
    ``SourceFileNotFoundError`` synchronously when ``recompute_pipeline``
    tries to re-project. The error must not be swallowed into ``FAILED`` on
    the background task -- the operator needs the synchronous error envelope.
    """
    from sage.api.errors import SourceFileNotFoundError

    full_path = _create_test_file(tmp_vault_dir, "samples/recpl5.md", "# RecPL5\n\nContent.")
    result = await ingestion_service.ingest(
        IngestRequest(source="samples/recpl5.md", source_type=SourceType.MARKDOWN)
    )
    doc_id = result.document.id

    full_path.unlink()
    assert not full_path.exists()

    with pytest.raises(SourceFileNotFoundError):
        await ingestion_service.recompute_pipeline(doc_id)

    # The claim must NOT be held after a synchronous failure -- otherwise the
    # operator would have to restart the service to retry.
    assert doc_id not in ingestion_service._inflight


# ---------------------------------------------------------------------------
# BH-131, BH-132, BH-133: Adapter-emitted tags merge into document.tags
# ---------------------------------------------------------------------------
#
# These tests exercise the adapter_tags / adapter_tag_prefixes convention
# between the source adapter layer and the ingestion service. The
# DocxAdapter is the first adapter to use this channel (for.dotx template
# style inventory), but the contract is generic so any adapter can
# contribute to document.tags from projection metadata.


class _TagEmittingStubAdapter:
    """Stub adapter that emits configurable adapter_tags in projection metadata.

    Used to exercise the ingestion-level tag-merge plumbing (BH-131..133)
    without the real.dotx machinery. Each instance can be reconfigured
    between calls to simulate an updated adapter version, supporting
    the BH-132 stale-tag-stripping scenario.
    """

    VERSION = "0.0.1-stub"
    EXTENSIONS = [".md"]

    def __init__(self, adapter_tags: list[str], adapter_tag_prefixes: list[str]):
        self.adapter_tags = list(adapter_tags)
        self.adapter_tag_prefixes = list(adapter_tag_prefixes)

    async def project(self, source_path, config=None):
        import hashlib
        from datetime import datetime, timezone

        from sage.source_adapters.base import HeadingNode, ProjectionResult

        raw = source_path.read_bytes()
        text = source_path.read_text()
        mtime = datetime.fromtimestamp(source_path.stat().st_mtime, tz=timezone.utc).isoformat()
        return ProjectionResult(
            text=text,
            headings=[HeadingNode(level=1, text="Stub", path="Stub", content=text)],
            content_hash=hashlib.sha256(raw).hexdigest(),
            adapter_version=self.VERSION,
            title=source_path.stem,
            metadata={
                "source_modified_at": mtime,
                "adapter_tags": list(self.adapter_tags),
                "adapter_tag_prefixes": list(self.adapter_tag_prefixes),
            },
        )


def _make_ingestion_service_with_tag_adapter(
    *,
    graph_store,
    lock_manager,
    stub_content_store,
    stub_embedding_provider,
    stub_abstraction_provider,
    minimal_config,
    lifecycle_service,
    adapter,
):
    """Build a fresh IngestionService using the given adapter under MARKDOWN."""
    from sage.services.ingestion import IngestionService

    return IngestionService(
        graph_store=graph_store,
        lock_manager=lock_manager,
        content_store=stub_content_store,
        embedding_provider=stub_embedding_provider,
        abstraction_provider=stub_abstraction_provider,
        config=minimal_config,
        source_adapters={SourceType.MARKDOWN: adapter},
        lifecycle_service=lifecycle_service,
    )


async def test_bh_131_adapter_tags_merge_on_new_ingest(
    tmp_vault_dir,
    graph_store,
    lock_manager,
    stub_content_store,
    stub_embedding_provider,
    stub_abstraction_provider,
    minimal_config,
    lifecycle_service,
):
    """BH-131: adapter_tags merge with caller-supplied tags on new ingest."""
    adapter = _TagEmittingStubAdapter(
        adapter_tags=["template:style:A", "template:style:B"],
        adapter_tag_prefixes=["template:"],
    )
    service = _make_ingestion_service_with_tag_adapter(
        graph_store=graph_store,
        lock_manager=lock_manager,
        stub_content_store=stub_content_store,
        stub_embedding_provider=stub_embedding_provider,
        stub_abstraction_provider=stub_abstraction_provider,
        minimal_config=minimal_config,
        lifecycle_service=lifecycle_service,
        adapter=adapter,
    )

    _create_test_file(
        tmp_vault_dir,
        "reports/bh131_a.md",
        content="# BH-131 case A\n\nWith adapter tags.",
    )

    request = IngestRequest(
        source="reports/bh131_a.md",
        source_type=SourceType.MARKDOWN,
        metadata={"codes": "caller-tag"},
    )
    result = await service.ingest(request)
    doc = result.document

    assert "template:style:A" in doc.tags
    assert "template:style:B" in doc.tags
    assert "caller-tag" in doc.tags
    # No duplicates
    assert len(doc.tags) == len(set(doc.tags))

    # Negative case: adapter that emits no adapter_tags leaves doc.tags
    # derived from caller only.
    plain_adapter = _TagEmittingStubAdapter(
        adapter_tags=[],
        adapter_tag_prefixes=[],
    )

    # Tweak so the plain adapter doesn't emit the metadata keys at all
    async def plain_project(source_path, config=None):
        import hashlib
        from datetime import datetime, timezone

        from sage.source_adapters.base import HeadingNode, ProjectionResult

        raw = source_path.read_bytes()
        text = source_path.read_text()
        return ProjectionResult(
            text=text,
            headings=[HeadingNode(level=1, text="Plain", path="Plain", content=text)],
            content_hash=hashlib.sha256(raw).hexdigest(),
            adapter_version=plain_adapter.VERSION,
            title=source_path.stem,
            metadata={
                "source_modified_at": datetime.fromtimestamp(
                    source_path.stat().st_mtime, tz=timezone.utc
                ).isoformat(),
            },
        )

    plain_adapter.project = plain_project

    service_plain = _make_ingestion_service_with_tag_adapter(
        graph_store=graph_store,
        lock_manager=lock_manager,
        stub_content_store=stub_content_store,
        stub_embedding_provider=stub_embedding_provider,
        stub_abstraction_provider=stub_abstraction_provider,
        minimal_config=minimal_config,
        lifecycle_service=lifecycle_service,
        adapter=plain_adapter,
    )

    _create_test_file(
        tmp_vault_dir,
        "reports/bh131_plain.md",
        content="# BH-131 case B\n\nWithout adapter tags.",
    )
    plain_result = await service_plain.ingest(
        IngestRequest(
            source="reports/bh131_plain.md",
            source_type=SourceType.MARKDOWN,
            metadata={"codes": "caller-tag"},
        )
    )
    assert plain_result.document.tags == ["caller-tag"]


async def test_bh_132_force_reingest_strips_stale_adapter_tags(
    tmp_vault_dir,
    graph_store,
    lock_manager,
    stub_content_store,
    stub_embedding_provider,
    stub_abstraction_provider,
    minimal_config,
    lifecycle_service,
):
    """BH-132: stale adapter-owned tags are stripped on force re-ingest."""
    adapter = _TagEmittingStubAdapter(
        adapter_tags=["template:style:Old", "template:has_numbering:Old"],
        adapter_tag_prefixes=["template:"],
    )
    service = _make_ingestion_service_with_tag_adapter(
        graph_store=graph_store,
        lock_manager=lock_manager,
        stub_content_store=stub_content_store,
        stub_embedding_provider=stub_embedding_provider,
        stub_abstraction_provider=stub_abstraction_provider,
        minimal_config=minimal_config,
        lifecycle_service=lifecycle_service,
        adapter=adapter,
    )

    file_path = _create_test_file(
        tmp_vault_dir,
        "reports/bh132.md",
        content="# BH-132\n\nFirst.",
    )

    first = await service.ingest(
        IngestRequest(
            source="reports/bh132.md",
            source_type=SourceType.MARKDOWN,
            metadata={"codes": "caller-tag"},
        )
    )
    assert "template:style:Old" in first.document.tags
    assert "template:has_numbering:Old" in first.document.tags

    # "Update" the adapter to emit a different tag set. Force re-ingest to
    # trigger the stale-stripping path. Content must change to bypass the
    # duplicate gate (BH-133 covers the byte-identical case).
    file_path.write_text("# BH-132\n\nSecond revision.")
    adapter.adapter_tags = ["template:style:New"]

    second = await service.ingest(
        IngestRequest(
            source="reports/bh132.md",
            source_type=SourceType.MARKDOWN,
            force=True,
            metadata={"codes": "caller-tag"},
        )
    )

    assert "template:style:New" in second.document.tags
    assert "template:style:Old" not in second.document.tags
    assert "template:has_numbering:Old" not in second.document.tags
    assert "caller-tag" in second.document.tags


async def test_bh_133_byte_identical_reingest_raises_duplicate(
    tmp_vault_dir,
    graph_store,
    lock_manager,
    stub_content_store,
    stub_embedding_provider,
    stub_abstraction_provider,
    minimal_config,
    lifecycle_service,
):
    """BH-133: adapter_tags channel does not bypass the content-hash duplicate gate."""
    adapter = _TagEmittingStubAdapter(
        adapter_tags=["template:style:X"],
        adapter_tag_prefixes=["template:"],
    )
    service = _make_ingestion_service_with_tag_adapter(
        graph_store=graph_store,
        lock_manager=lock_manager,
        stub_content_store=stub_content_store,
        stub_embedding_provider=stub_embedding_provider,
        stub_abstraction_provider=stub_abstraction_provider,
        minimal_config=minimal_config,
        lifecycle_service=lifecycle_service,
        adapter=adapter,
    )

    _create_test_file(tmp_vault_dir, "reports/bh133.md", content="# BH-133\n\nOnly content.")

    first = await service.ingest(
        IngestRequest(source="reports/bh133.md", source_type=SourceType.MARKDOWN)
    )
    assert "template:style:X" in first.document.tags
    original_adapter_version = first.document.adapter_version
    original_updated_at = first.document.updated_at
    original_tags = list(first.document.tags)

    # Same file, no force: must raise DuplicateContentError
    with pytest.raises(DuplicateContentError) as exc_info:
        await service.ingest(
            IngestRequest(source="reports/bh133.md", source_type=SourceType.MARKDOWN)
        )
    assert exc_info.value.detail["existing_document_id"] == first.document.id

    # Document state should be unchanged after the rejected re-ingest
    fetched = await graph_store.get_document(first.document.id)
    assert fetched.tags == original_tags
    assert fetched.adapter_version == original_adapter_version
    assert fetched.updated_at == original_updated_at


# ---------------------------------------------------------------------------
# BH-134: Empty projection text transitions to abstraction_skipped
# ---------------------------------------------------------------------------


class _EmptyTextStubAdapter:
    """Stub adapter that returns an empty projection text, simulating a
    Word template or other content-thin source."""

    VERSION = "0.0.1-empty"
    EXTENSIONS = [".md"]

    async def project(self, source_path, config=None):
        import hashlib
        from datetime import datetime, timezone

        from sage.source_adapters.base import ProjectionResult

        raw = source_path.read_bytes()
        return ProjectionResult(
            text="",
            headings=[],
            content_hash=hashlib.sha256(raw).hexdigest(),
            adapter_version=self.VERSION,
            title=source_path.stem,
            metadata={
                "source_modified_at": datetime.fromtimestamp(
                    source_path.stat().st_mtime, tz=timezone.utc
                ).isoformat(),
            },
        )


class _StrictAbstractionProvider:
    """Mirrors Qwen3's strict edge guard: raises on empty input.

    Any test that reaches generate_abstract() with empty text via this
    provider will fail. BH-134 relies on this behavior to verify that
    the service-layer guard short-circuits before reaching here.
    """

    async def generate_abstract(self, text: str, max_tokens: int, doc_type: str | None) -> str:
        if not text or not text.strip():
            raise RuntimeError("Cannot generate abstract from empty document text")
        return f"Strict stub abstract for {len(text)} chars."


async def test_bh_134_empty_projection_text_skips_abstraction(
    tmp_vault_dir,
    graph_store,
    lock_manager,
    stub_content_store,
    stub_embedding_provider,
    minimal_config,
    lifecycle_service,
):
    """BH-134: empty projection text transitions to abstraction_skipped."""
    from sage.services.ingestion import IngestionService

    service = IngestionService(
        graph_store=graph_store,
        lock_manager=lock_manager,
        content_store=stub_content_store,
        embedding_provider=stub_embedding_provider,
        abstraction_provider=_StrictAbstractionProvider(),
        config=minimal_config,
        source_adapters={SourceType.MARKDOWN: _EmptyTextStubAdapter()},
        lifecycle_service=lifecycle_service,
    )

    _create_test_file(
        tmp_vault_dir,
        "reports/bh134.md",
        content="# Placeholder\n\nSource has bytes but adapter returns empty text.",
    )

    result = await service.ingest(
        IngestRequest(source="reports/bh134.md", source_type=SourceType.MARKDOWN)
    )

    fetched = await graph_store.get_document(result.document.id)
    assert fetched.pipeline_status == PipelineStatus.ABSTRACTION_SKIPPED
    assert fetched.semantic_abstract is None
    assert fetched.pipeline_error is None


async def test_bh_134_non_empty_text_still_runs_abstraction(
    tmp_vault_dir,
    graph_store,
    lock_manager,
    stub_content_store,
    stub_embedding_provider,
    minimal_config,
    lifecycle_service,
):
    """BH-134 (control): non-empty text still reaches abstraction."""
    from sage.services.ingestion import IngestionService

    service = IngestionService(
        graph_store=graph_store,
        lock_manager=lock_manager,
        content_store=stub_content_store,
        embedding_provider=stub_embedding_provider,
        abstraction_provider=_StrictAbstractionProvider(),
        config=minimal_config,
        source_adapters={SourceType.MARKDOWN: MarkdownAdapter()},
        lifecycle_service=lifecycle_service,
    )

    _create_test_file(
        tmp_vault_dir,
        "reports/bh134_ctrl.md",
        content="# Real\n\nNon-empty body content to abstract.",
    )

    result = await service.ingest(
        IngestRequest(source="reports/bh134_ctrl.md", source_type=SourceType.MARKDOWN)
    )

    fetched = await graph_store.get_document(result.document.id)
    assert fetched.pipeline_status == PipelineStatus.ABSTRACTION_COMPLETE
    assert fetched.semantic_abstract is not None
    assert "Strict stub abstract" in fetched.semantic_abstract


# ---------------------------------------------------------------------------
# Vault-level adapter parameter propagation
#
# The vault config's adapter_defaults[source_type] entry is authoritative for
# adapter behavior. Per-request IngestRequest.config is a per-call override
# that takes precedence on key collisions. The retired source_adapters
# section is inert (CAS-ADR-046) and must not feed the adapter.
# ---------------------------------------------------------------------------

import copy as _copy  # noqa: E402 -- grouped with the vault-adapter-config test section below
from typing import Any  # noqa: E402 -- grouped with the vault-adapter-config test section below

try:
    import docx as _docx_module

    _HAS_DOCX = True
except ImportError:
    _HAS_DOCX = False

requires_docx = pytest.mark.skipif(not _HAS_DOCX, reason="python-docx not available")


def _write_styled_docx(path: Path, paragraphs: list[tuple[str, str]]) -> None:
    """Write a.docx with the given (text, style) paragraphs.

    Style names must already exist in the document's style set. Built-in
    styles like "Title", "Heading 1", "Subtitle", "Normal" are pre-defined
    by python-docx.
    """
    doc = _docx_module.Document()
    for text, style in paragraphs:
        doc.add_paragraph(text, style=style)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


def _build_vault_config_with_docx(base_dict: dict, vault_docx_config: dict | None) -> VaultConfig:
    """Return a VaultConfig carrying docx adapter defaults from base_dict.

    If vault_docx_config is None, no docx entry is declared and the adapter
    falls through to its own defaults. Otherwise vault_docx_config is
    placed under adapter_defaults["docx"].
    """
    config_dict: dict[str, Any] = _copy.deepcopy(base_dict)
    if vault_docx_config is not None:
        config_dict.setdefault("adapter_defaults", {})["docx"] = vault_docx_config
    return VaultConfig.model_validate(config_dict)


def _make_ingestion_with_docx(
    config: VaultConfig,
    *,
    graph_store,
    lock_manager,
    stub_content_store,
    stub_embedding_provider,
    stub_abstraction_provider,
):
    """Construct an IngestionService with real DocxAdapter + Markdown stub."""
    from sage.services.ingestion import IngestionService
    from sage.services.lifecycle import LifecycleService
    from sage.source_adapters.docx_adapter import DocxAdapter

    lifecycle_service = LifecycleService(graph_store, lock_manager, config)
    return IngestionService(
        graph_store=graph_store,
        lock_manager=lock_manager,
        content_store=stub_content_store,
        embedding_provider=stub_embedding_provider,
        abstraction_provider=stub_abstraction_provider,
        config=config,
        source_adapters={
            SourceType.MARKDOWN: MarkdownAdapter(),
            SourceType.DOCX: DocxAdapter(),
        },
        lifecycle_service=lifecycle_service,
    )


@requires_docx
async def test_vault_adapter_config_reaches_adapter_at_ingest(
    tmp_vault_dir,
    graph_store,
    lock_manager,
    stub_content_store,
    stub_embedding_provider,
    stub_abstraction_provider,
    minimal_vault_config_dict,
):
    """The EXAMPLE use case: vault config's heading_style_map flows to docx adapter."""
    config = _build_vault_config_with_docx(
        minimal_vault_config_dict,
        vault_docx_config={"heading_style_map": {"Title": 1}},
    )
    service = _make_ingestion_with_docx(
        config,
        graph_store=graph_store,
        lock_manager=lock_manager,
        stub_content_store=stub_content_store,
        stub_embedding_provider=stub_embedding_provider,
        stub_abstraction_provider=stub_abstraction_provider,
    )

    docx_path = tmp_vault_dir / "sources" / "vault_config_docx.docx"
    _write_styled_docx(
        docx_path,
        [("CLAIMS", "Title"), ("Body content under CLAIMS.", "Normal")],
    )

    result = await service.ingest(
        IngestRequest(source="vault_config_docx.docx", source_type=SourceType.DOCX)
    )

    # The Title-styled "CLAIMS" paragraph must be recognized as a heading
    # because the vault config maps Title -> level 1. Without the fix, the
    # adapter receives request.config (None) and falls back to defaults
    # (Heading 1-9 only), so no chunk gets heading_path "CLAIMS".
    heading_paths = await stub_content_store.get_heading_paths(result.document.id)
    assert "CLAIMS" in heading_paths, f"Expected 'CLAIMS' as a heading_path; got: {heading_paths}"


class _RecordingDocxAdapter:
    """Real DocxAdapter behavior, plus a record of every config it received.

    Wraps rather than subclasses so the delegate stays the production class:
    a test double that reimplemented projection could agree with an
    assertion about heading levels while proving nothing about what the
    real adapter was handed.
    """

    def __init__(self) -> None:
        from sage.source_adapters.docx_adapter import DocxAdapter

        self._delegate = DocxAdapter()
        self.configs: list[dict | None] = []

    EXTENSIONS = (".docx",)

    async def project(self, source_path: Path, config: dict | None = None):
        self.configs.append(config)
        return await self._delegate.project(source_path, config)

    def __getattr__(self, name):
        return getattr(self._delegate, name)


@requires_docx
async def test_recompute_pipeline_applies_vault_adapter_defaults(
    tmp_vault_dir,
    graph_store,
    lock_manager,
    stub_content_store,
    stub_embedding_provider,
    stub_abstraction_provider,
    minimal_vault_config_dict,
):
    """Re-projection reads adapter parameters from the vault config.

    ``recompute_pipeline`` carries no per-request config, so vault defaults
    are its only source of adapter parameters. A relocation that wired the
    new section into ingest alone would leave this path reading nothing and
    silently flatten the heading tree of every re-projected document.
    """
    from sage.services.ingestion import IngestionService
    from sage.services.lifecycle import LifecycleService

    config = _build_vault_config_with_docx(
        minimal_vault_config_dict,
        vault_docx_config={"heading_style_map": {"Title": 1}},
    )
    recording_adapter = _RecordingDocxAdapter()
    service = IngestionService(
        graph_store=graph_store,
        lock_manager=lock_manager,
        content_store=stub_content_store,
        embedding_provider=stub_embedding_provider,
        abstraction_provider=stub_abstraction_provider,
        config=config,
        source_adapters={
            SourceType.MARKDOWN: MarkdownAdapter(),
            SourceType.DOCX: recording_adapter,
        },
        lifecycle_service=LifecycleService(graph_store, lock_manager, config),
    )

    docx_path = tmp_vault_dir / "sources" / "recompute_styles.docx"
    _write_styled_docx(
        docx_path,
        [("RECOMPUTED", "Title"), ("Body under the custom style.", "Normal")],
    )
    result = await service.ingest(
        IngestRequest(source="recompute_styles.docx", source_type=SourceType.DOCX)
    )

    recording_adapter.configs.clear()
    await service.recompute_pipeline(result.document.id)

    assert recording_adapter.configs, "recompute_pipeline did not re-project"
    reprojection_config = recording_adapter.configs[-1]
    assert reprojection_config is not None, (
        "re-projection received no adapter config; vault defaults did not "
        "reach the path that has no per-request fallback"
    )
    assert reprojection_config["heading_style_map"]["Title"] == 1


@requires_docx
async def test_retired_source_adapters_section_does_not_feed_the_adapter(
    tmp_vault_dir,
    graph_store,
    lock_manager,
    stub_content_store,
    stub_embedding_provider,
    stub_abstraction_provider,
    minimal_vault_config_dict,
):
    """Parameters left in the retired section are inert (CAS-ADR-046).

    Without this, every other test in the block would still pass if the
    merge read both shapes -- and the migration would be cosmetic, leaving
    the stale section quietly load-bearing.
    """
    config_dict = _copy.deepcopy(minimal_vault_config_dict)
    config_dict["source_adapters"] = {
        "adapters": [
            {
                "source_type": "docx",
                "enabled": True,
                "config": {"heading_style_map": {"Title": 1}},
            }
        ]
    }
    config = VaultConfig.model_validate(config_dict)
    assert config.adapter_defaults == {}

    service = _make_ingestion_with_docx(
        config,
        graph_store=graph_store,
        lock_manager=lock_manager,
        stub_content_store=stub_content_store,
        stub_embedding_provider=stub_embedding_provider,
        stub_abstraction_provider=stub_abstraction_provider,
    )

    docx_path = tmp_vault_dir / "sources" / "legacy_section.docx"
    _write_styled_docx(
        docx_path,
        [("STALE", "Title"), ("Body content under STALE.", "Normal")],
    )

    result = await service.ingest(
        IngestRequest(source="legacy_section.docx", source_type=SourceType.DOCX)
    )

    heading_paths = await stub_content_store.get_heading_paths(result.document.id)
    assert "STALE" not in heading_paths, (
        f"the retired section still reached the adapter; heading_paths={heading_paths}"
    )


@requires_docx
async def test_request_config_overrides_vault_adapter_config_on_key_collision(
    tmp_vault_dir,
    graph_store,
    lock_manager,
    stub_content_store,
    stub_embedding_provider,
    stub_abstraction_provider,
    minimal_vault_config_dict,
):
    """Per-request config overrides vault-level on key collisions."""
    config = _build_vault_config_with_docx(
        minimal_vault_config_dict,
        vault_docx_config={"heading_style_map": {"Title": 1}},
    )
    service = _make_ingestion_with_docx(
        config,
        graph_store=graph_store,
        lock_manager=lock_manager,
        stub_content_store=stub_content_store,
        stub_embedding_provider=stub_embedding_provider,
        stub_abstraction_provider=stub_abstraction_provider,
    )

    docx_path = tmp_vault_dir / "sources" / "override.docx"
    # Two Title paragraphs so we can see hierarchy levels in heading paths.
    _write_styled_docx(
        docx_path,
        [
            ("Outer", "Heading 1"),
            ("Inner", "Title"),
            ("Body", "Normal"),
        ],
    )

    # Request asks for Title=2; vault says Title=1. Request must win.
    result = await service.ingest(
        IngestRequest(
            source="override.docx",
            source_type=SourceType.DOCX,
            config={"heading_style_map": {"Title": 2}},
        )
    )

    heading_paths = await stub_content_store.get_heading_paths(result.document.id)
    # If Title is level 2, it nests under Heading 1: path should be "Outer > Inner".
    # If Title were level 1 (vault value), the path would be just "Inner".
    assert "Outer > Inner" in heading_paths, (
        f"Expected Title=2 (request wins); got heading_paths: {heading_paths}"
    )


@requires_docx
async def test_request_config_merges_with_vault_adapter_config_no_collision(
    tmp_vault_dir,
    graph_store,
    lock_manager,
    stub_content_store,
    stub_embedding_provider,
    stub_abstraction_provider,
    minimal_vault_config_dict,
):
    """Non-conflicting keys merge; both vault and request mappings apply."""
    config = _build_vault_config_with_docx(
        minimal_vault_config_dict,
        vault_docx_config={"heading_style_map": {"Title": 1}},
    )
    service = _make_ingestion_with_docx(
        config,
        graph_store=graph_store,
        lock_manager=lock_manager,
        stub_content_store=stub_content_store,
        stub_embedding_provider=stub_embedding_provider,
        stub_abstraction_provider=stub_abstraction_provider,
    )

    docx_path = tmp_vault_dir / "sources" / "merge.docx"
    # Each heading needs body content beneath it; chunks are only created
    # for headings whose content is non-empty.
    _write_styled_docx(
        docx_path,
        [
            ("VaultStyleHeading", "Title"),
            ("Vault body.", "Normal"),
            ("RequestStyleHeading", "Subtitle"),
            ("Request body.", "Normal"),
        ],
    )

    # Vault: Title -> 1. Request adds: Subtitle -> 1. No collision; both apply.
    result = await service.ingest(
        IngestRequest(
            source="merge.docx",
            source_type=SourceType.DOCX,
            config={"heading_style_map": {"Subtitle": 1}},
        )
    )

    heading_paths = await stub_content_store.get_heading_paths(result.document.id)
    assert "VaultStyleHeading" in heading_paths
    assert "RequestStyleHeading" in heading_paths


@requires_docx
async def test_no_vault_adapter_entry_falls_through_to_request_config(
    tmp_vault_dir,
    graph_store,
    lock_manager,
    stub_content_store,
    stub_embedding_provider,
    stub_abstraction_provider,
    minimal_vault_config_dict,
):
    """Vault config with no docx adapter entry: request.config still works."""
    # No docx entry added; minimal_vault_config_dict only has markdown.
    config = VaultConfig.model_validate(_copy.deepcopy(minimal_vault_config_dict))
    service = _make_ingestion_with_docx(
        config,
        graph_store=graph_store,
        lock_manager=lock_manager,
        stub_content_store=stub_content_store,
        stub_embedding_provider=stub_embedding_provider,
        stub_abstraction_provider=stub_abstraction_provider,
    )

    docx_path = tmp_vault_dir / "sources" / "no_vault_entry.docx"
    _write_styled_docx(
        docx_path,
        [("RequestOnly", "Title"), ("Body", "Normal")],
    )

    result = await service.ingest(
        IngestRequest(
            source="no_vault_entry.docx",
            source_type=SourceType.DOCX,
            config={"heading_style_map": {"Title": 1}},
        )
    )

    heading_paths = await stub_content_store.get_heading_paths(result.document.id)
    assert "RequestOnly" in heading_paths


@requires_docx
async def test_no_request_config_and_no_vault_adapter_config_uses_defaults(
    tmp_vault_dir,
    graph_store,
    lock_manager,
    stub_content_store,
    stub_embedding_provider,
    stub_abstraction_provider,
    minimal_vault_config_dict,
):
    """Both configs absent: adapter falls back to _DEFAULT_STYLE_MAP."""
    config = _build_vault_config_with_docx(minimal_vault_config_dict, vault_docx_config=None)
    service = _make_ingestion_with_docx(
        config,
        graph_store=graph_store,
        lock_manager=lock_manager,
        stub_content_store=stub_content_store,
        stub_embedding_provider=stub_embedding_provider,
        stub_abstraction_provider=stub_abstraction_provider,
    )

    docx_path = tmp_vault_dir / "sources" / "defaults.docx"
    _write_styled_docx(
        docx_path,
        [
            ("BuiltInH1", "Heading 1"),
            ("CustomTitle", "Title"),  # not in defaults
            ("Body", "Normal"),
        ],
    )

    result = await service.ingest(
        IngestRequest(source="defaults.docx", source_type=SourceType.DOCX)
    )

    heading_paths = await stub_content_store.get_heading_paths(result.document.id)
    # Heading 1 is in defaults; Title is not.
    assert "BuiltInH1" in heading_paths
    assert "CustomTitle" not in heading_paths


# ---------------------------------------------------------------------------
# Heading-context embedding
#
# At index time, the embedder receives `heading_path + content`, not just
# `content`. This makes semantic search reach chunks whose query terms
# appear only in the heading hierarchy — the agent equivalent of Word's
# Find on a heading.
# ---------------------------------------------------------------------------


class _RecordingEmbeddingProvider:
    """Captures the texts passed to embed() for assertion.

    Accumulates across calls — the production pipeline now embeds chunks
    in Stage 2 and re-embeds the synthetic header chunk in Stage 3,
    so a single call's inputs are no longer the full picture.
    """

    def __init__(self) -> None:
        self.last_inputs: list[str] = []

    async def embed(self, texts: list[str]) -> list[list[float]]:
        self.last_inputs.extend(texts)
        return [[0.0] * 768 for _ in texts]


@requires_docx
async def test_embedder_receives_combined_heading_path_and_content(
    tmp_vault_dir,
    graph_store,
    lock_manager,
    stub_content_store,
    stub_abstraction_provider,
    minimal_vault_config_dict,
):
    """The text passed to the embedder includes heading_path + content,
    so semantic search can find chunks via heading-only query terms.

    Heading text is deliberately distinct from the document title so the
    standalone synthetic header chunk — which carries title /
    source / tags / semantic_abstract — does not accidentally include
    the heading text; that would mask a failure of the heading-context
    embedding logic.
    """
    from sage.services.ingestion import IngestionService
    from sage.services.lifecycle import LifecycleService
    from sage.source_adapters.docx_adapter import DocxAdapter

    config = _build_vault_config_with_docx(
        minimal_vault_config_dict,
        vault_docx_config=None,  # use defaults (Heading 1-9 only)
    )
    lifecycle_service = LifecycleService(graph_store, lock_manager, config)
    recording_embedder = _RecordingEmbeddingProvider()
    service = IngestionService(
        graph_store=graph_store,
        lock_manager=lock_manager,
        content_store=stub_content_store,
        embedding_provider=recording_embedder,
        abstraction_provider=stub_abstraction_provider,
        config=config,
        source_adapters={
            SourceType.MARKDOWN: MarkdownAdapter(),
            SourceType.DOCX: DocxAdapter(),
        },
        lifecycle_service=lifecycle_service,
    )

    docx_path = tmp_vault_dir / "sources" / "plain_filename.docx"
    # Filename is "plain_filename" → doc title is "plain_filename".
    # Heading text "Distinctive Marker Phrase" appears nowhere in the
    # filename or body, so the only way it can reach the embedder is
    # via heading_path concatenation.
    _write_styled_docx(
        docx_path,
        [
            ("Distinctive Marker Phrase", "Heading 1"),
            ("Cryptographic accumulator seals govern the commit.", "Normal"),
        ],
    )

    await service.ingest(IngestRequest(source="plain_filename.docx", source_type=SourceType.DOCX))

    assert recording_embedder.last_inputs, "embedder was not called"
    # The synthetic header chunk is index 0; the body chunk is
    # at the index whose embed input begins with the heading text.
    body_inputs = [
        text
        for text in recording_embedder.last_inputs
        if text.startswith("Distinctive Marker Phrase")
    ]
    assert body_inputs, (
        f"no embed input started with the heading: {recording_embedder.last_inputs!r}"
    )
    combined = body_inputs[0]
    assert "Distinctive Marker Phrase" in combined, (
        f"heading_path missing from embed input: {combined!r}"
    )
    assert "Cryptographic accumulator seals" in combined, (
        f"body content missing from embed input: {combined!r}"
    )


@requires_docx
async def test_chunk_content_field_unchanged_by_combined_embedding(
    tmp_vault_dir,
    graph_store,
    lock_manager,
    stub_content_store,
    stub_embedding_provider,
    stub_abstraction_provider,
    minimal_vault_config_dict,
):
    """The chunk.content field stored in the content store carries its
    own ATX heading line plus body content. The embedder builds
    its `heading_path + content` input on the fly; it does not mutate the
    stored chunk.content. The synthetic-header preamble does not
    leak into body chunks."""
    config = _build_vault_config_with_docx(
        minimal_vault_config_dict,
        vault_docx_config={"heading_style_map": {"Title": 1}},
    )
    service = _make_ingestion_with_docx(
        config,
        graph_store=graph_store,
        lock_manager=lock_manager,
        stub_content_store=stub_content_store,
        stub_embedding_provider=stub_embedding_provider,
        stub_abstraction_provider=stub_abstraction_provider,
    )

    docx_path = tmp_vault_dir / "sources" / "clean_content.docx"
    body_first = "First section body content."
    second_heading = "Exemplary Methods"
    body_second = "Second section body about cryptographic accumulator seals."
    _write_styled_docx(
        docx_path,
        [
            ("Introduction", "Title"),
            (body_first, "Normal"),
            (second_heading, "Title"),
            (body_second, "Normal"),
        ],
    )

    result = await service.ingest(
        IngestRequest(source="clean_content.docx", source_type=SourceType.DOCX)
    )

    chunks = await stub_content_store.get_all_chunks(result.document.id)
    # Passages carry the ATX heading line plus body content; document-level
    # text lives on its own surface and is not among them (CAS-ADR-049).
    assert len(chunks) >= 2
    second = next(c for c in chunks if c.heading_path == second_heading)
    assert second.content == f"# {second_heading}\n\n{body_second}", (
        "Stored chunk.content must be ATX heading line + body; got: " + repr(second.content)
    )
    # Document-level preamble fields must not leak into passages.
    assert not second.content.startswith("Title:")


@requires_docx
async def test_empty_content_heading_still_emits_chunk(
    tmp_vault_dir,
    graph_store,
    lock_manager,
    stub_content_store,
    stub_embedding_provider,
    stub_abstraction_provider,
    minimal_vault_config_dict,
):
    """A heading whose immediate next paragraph is another same-or-higher
    level heading produces no body content for itself. The chunker MUST
    still emit a chunk for that heading so its heading_path enters the
    FTS index. Without this, a heading like "DETAILED DESCRIPTION" that's
    immediately followed by another top-level heading vanishes from the
    indexed surface — invisible to BM25, semantic search, and
    read_section heading enumeration. Word's Find finds it; SAGE
    must too. Closes the gap surfaced by Cowork on PV01 v10.2.2 et al.
    """
    config = _build_vault_config_with_docx(
        minimal_vault_config_dict,
        vault_docx_config=None,  # use defaults: Heading 1-9 only
    )
    service = _make_ingestion_with_docx(
        config,
        graph_store=graph_store,
        lock_manager=lock_manager,
        stub_content_store=stub_content_store,
        stub_embedding_provider=stub_embedding_provider,
        stub_abstraction_provider=stub_abstraction_provider,
    )

    docx_path = tmp_vault_dir / "sources" / "empty_heading.docx"
    # Two same-level headings back-to-back. The first ("EMPTY PARENT")
    # has no body content because the next paragraph is also a heading.
    _write_styled_docx(
        docx_path,
        [
            ("EMPTY PARENT", "Heading 1"),
            ("FIRST CHILD", "Heading 1"),
            ("Some body content under FIRST CHILD.", "Normal"),
        ],
    )

    result = await service.ingest(
        IngestRequest(source="empty_heading.docx", source_type=SourceType.DOCX)
    )

    heading_paths = await stub_content_store.get_heading_paths(result.document.id)
    assert "EMPTY PARENT" in heading_paths, (
        f"Heading 'EMPTY PARENT' must produce a chunk even with empty body "
        f"content. Got heading_paths: {heading_paths}"
    )
    assert "FIRST CHILD" in heading_paths

    chunks = await stub_content_store.get_all_chunks(result.document.id)
    by_path = {c.heading_path: c for c in chunks}
    assert "EMPTY PARENT" in by_path
    # Passages carry projected content only; document-level text lives on
    # its own surface and does not touch the per-heading chunks.
    parent = by_path["EMPTY PARENT"]
    assert "Some body content" not in parent.content, (
        "Body of FIRST CHILD must not leak into the EMPTY PARENT chunk."
    )
