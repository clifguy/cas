"""MCP profile-invariance tests (MPI-001 through MPI-016).

Prove the MCP document import/export byte channel behaves identically
whether the vault's source store is bound to the local filesystem or to a
cloud document store: same call shapes, same results, caller-local paths
throughout. Specifications with per-test anti-coincidental-pass rationale
live in ``mcp_profile_invariance_tests.md``; every test here runs once per
binding via the ``vault_source_backend`` fixture.

Follows the direct-call MCP test pattern (tool functions invoked with a
pre-initialized vault registry, bypassing transport).
"""

import asyncio
import hashlib
import inspect
import json
from pathlib import Path

import pytest

import sage.mcp_server as _mcp
from sage.adapters.stubs import (
    StubAbstractionProvider,
    StubContentStore,
    StubEmbeddingProvider,
)
from sage.config import VaultConfig
from sage.mcp_server import (
    bulk_ingest_document,
    get_document,
    ingest_document,
    list_directory,
    read_projection,
)
from tests.sage.conftest import initialize_services_for_test

_VAULT_ID = "test_vault"


def _parse(result: str | dict) -> dict:
    if isinstance(result, dict):
        return result
    return json.loads(result)


def _sha256(data: bytes) -> str:
    return "sha256:" + hashlib.sha256(data).hexdigest()


@pytest.fixture
async def mpi_vault(minimal_vault_config_dict, vault_source_backend):
    """One registered vault running over the parameterized source binding."""
    config = VaultConfig.model_validate(minimal_vault_config_dict)
    async with initialize_services_for_test(
        config,
        content_store=StubContentStore(),
        embedding_provider=StubEmbeddingProvider(),
        abstraction_provider=StubAbstractionProvider(),
    ) as services:
        _mcp._vaults[_VAULT_ID] = services
        try:
            yield services, config, vault_source_backend
        finally:
            # Drain the fire-and-forget pipeline tasks ingest dispatches
            # before services tear down (same discipline as the other
            # ingest-heavy MCP fixtures).
            await asyncio.sleep(0.3)
            _mcp._vaults.pop(_VAULT_ID, None)


async def _ingest_local_file(tmp_path: Path, name: str, body: str) -> tuple[Path, dict]:
    """Write a caller-local file (outside storage_root) and ingest it."""
    inbox = tmp_path / "caller_inbox"
    inbox.mkdir(exist_ok=True)
    src = inbox / name
    src.write_text(body)
    result = _parse(await ingest_document(_VAULT_ID, str(src), "markdown"))
    return src, result


try:
    import docx as _docx

    _HAS_DOCX = True
except ImportError:  # pragma: no cover -- exercised only without python-docx
    _HAS_DOCX = False

requires_docx = pytest.mark.skipif(not _HAS_DOCX, reason="python-docx not available")


def _write_office_file(tmp_path: Path, name: str, heading: str) -> Path:
    """Write a caller-local ``.docx`` (outside storage_root) and return its path.

    An Office package, not markdown, because the retaining store rewrites Office
    packages at rest: this is the format on which the delivered bytes and the
    stored bytes actually diverge.
    """
    inbox = tmp_path / "caller_inbox"
    inbox.mkdir(exist_ok=True)
    src = inbox / name
    document = _docx.Document()
    document.add_heading(heading, level=1)
    document.add_paragraph("Body of the office probe.")
    document.save(str(src))
    return src


async def test_mpi_001_ingest_caller_local_absolute_path(mpi_vault, tmp_path):
    """MPI-001: an absolute caller-local source ingests under either binding,
    retained byte-identical on the active backend."""
    _services, config, handle = mpi_vault
    src, result = await _ingest_local_file(tmp_path, "mpi_note.md", "# MPI\n\nBody one.")

    assert "error" not in result, result
    assert "id" in result
    assert result["source_path"] == "imports/mpi_note.md"

    retained = handle.retained_bytes(config.vault.storage_root, "imports/mpi_note.md")
    assert retained == src.read_bytes()
    if handle.fake_client is not None:
        assert handle.fake_client.source_uploads == 1


async def test_mpi_002_bulk_ingest_caller_local_files(mpi_vault, tmp_path):
    """MPI-002: bulk ingest of caller-local files succeeds under either
    binding with an identical summary shape."""
    _services, config, handle = mpi_vault
    inbox = tmp_path / "caller_inbox"
    inbox.mkdir()
    one = inbox / "alpha_note.md"
    one.write_text("# Alpha\n\nFirst body.")
    two = inbox / "beta_note.md"
    two.write_text("# Beta\n\nSecond body.")

    result = _parse(
        await bulk_ingest_document(
            _VAULT_ID,
            [
                {"file_path": str(one), "source_type": "markdown"},
                {"file_path": str(two), "source_type": "markdown"},
            ],
        )
    )

    assert result.get("error_count") == 0, result
    assert result["documents_created"]["new"] == 2
    for key in ("documents_created", "edges_created", "edges_staged", "errors"):
        assert key in result
    for src in (one, two):
        retained = handle.retained_bytes(config.vault.storage_root, f"imports/{src.name}")
        assert retained == src.read_bytes()


async def test_mpi_003_list_directory_discovery_and_vault_check(mpi_vault, tmp_path):
    """MPI-003: discovery walks the caller-local directory; the
    already-in-vault check resolves through the graph store, so a
    previously ingested file reports ``unchanged`` under either binding."""
    _services, _config, _handle = mpi_vault
    scan_dir = tmp_path / "caller_scan"
    scan_dir.mkdir()
    md = scan_dir / "gamma_note.md"
    md.write_text("# Gamma\n\nScan body.")
    (scan_dir / "opaque.xyz").write_text("no adapter")

    first = _parse(await list_directory(_VAULT_ID, str(scan_dir)))
    assert set(first) == {"files", "warnings", "truncated"}
    by_name = {Path(f["file_path"]).name: f for f in first["files"]}
    assert by_name["gamma_note.md"]["sage_status"] == "new"
    assert by_name["opaque.xyz"]["sage_status"] == "no_adapter"

    ingest = _parse(await ingest_document(_VAULT_ID, str(md), "markdown"))
    assert "error" not in ingest, ingest

    second = _parse(await list_directory(_VAULT_ID, str(scan_dir)))
    by_name = {Path(f["file_path"]).name: f for f in second["files"]}
    assert by_name["gamma_note.md"]["sage_status"] == "unchanged"


async def test_mpi_004_get_document_write_to_path_round_trip(mpi_vault, tmp_path):
    """MPI-004: write_to_path delivers the retained source bytes to a
    caller-local target under either binding."""
    _services, _config, _handle = mpi_vault
    src, ingest = await _ingest_local_file(tmp_path, "delta_note.md", "# Delta\n\nRound trip.")
    target = tmp_path / "caller_out" / "delta_export.md"
    target.parent.mkdir()

    result = _parse(await get_document(_VAULT_ID, ingest["id"], write_to_path=str(target)))

    assert "error" not in result, result
    original = src.read_bytes()
    assert target.read_bytes() == original
    assert result["written_to"] == str(target)
    assert result["content_size"] == len(original)
    assert result["content_hash"] == _sha256(original)
    assert result.get("content") is None
    assert result["read_meta"]["body_present"] is False


async def test_mpi_005_read_projection_spill_skips_vault_source_store(mpi_vault, tmp_path):
    """MPI-005: projection spill writes caller-locally from the content
    store; the vault-source store is never consulted."""
    _services, _config, handle = mpi_vault
    _src, ingest = await _ingest_local_file(tmp_path, "epsilon_note.md", "# Epsilon\n\nSpill body.")
    await asyncio.sleep(0.5)
    target = tmp_path / "caller_out" / "epsilon_projection.md"
    target.parent.mkdir()

    if handle.fake_client is not None:
        reads_before = handle.fake_client.source_reads
        streams_before = len(handle.fake_client.source_streams)
        stats_before = handle.fake_client.source_stats

    result = _parse(await read_projection(_VAULT_ID, ingest["id"], write_to_path=str(target)))

    assert "error" not in result, result
    assert result["written_to"] == str(target)
    assert result.get("projection_text") is None
    assert result["content_size"] > 0
    spilled = target.read_text()
    assert "Spill body." in spilled
    if handle.fake_client is not None:
        assert handle.fake_client.source_reads == reads_before
        assert len(handle.fake_client.source_streams) == streams_before
        assert handle.fake_client.source_stats == stats_before


async def test_mpi_006_export_channel_streams_above_inline_ceiling(
    mpi_vault, tmp_path, monkeypatch
):
    """MPI-006: with the inline ceiling pinned tiny, inline delivery refuses
    while write_to_path round-trips the full payload through the streaming
    read under either binding."""
    _services, _config, handle = mpi_vault
    body = "# Big\n\n" + ("payload line\n" * 5000)  # ~64 KiB
    src, ingest = await _ingest_local_file(tmp_path, "zeta_big.md", body)
    monkeypatch.setenv("SAGE_MAX_INLINE_CONTENT_BYTES", "1024")

    inline = _parse(await get_document(_VAULT_ID, ingest["id"], include_content=True))
    assert inline["error"] == "content_too_large"

    if handle.fake_client is not None:
        reads_before = handle.fake_client.source_reads

    target = tmp_path / "caller_out" / "zeta_export.md"
    target.parent.mkdir()
    spilled = _parse(await get_document(_VAULT_ID, ingest["id"], write_to_path=str(target)))

    assert "error" not in spilled, spilled
    original = src.read_bytes()
    assert len(original) > 1024
    assert target.read_bytes() == original
    assert spilled["content_hash"] == _sha256(original)
    if handle.fake_client is not None:
        assert handle.fake_client.source_streams, "delivery must use the streaming read"
        assert handle.fake_client.source_reads == reads_before, (
            "write_to_path delivery must not fall back to the buffered whole-body read"
        )


@requires_docx
async def test_mpi_009_provenance_hash_is_the_delivered_hash(mpi_vault, tmp_path):
    """MPI-009: the recorded provenance hash is the SHA-256 of the bytes the
    caller delivered, under either binding, even for a format the retaining
    store rewrites at rest.

    Anti-coincidental-pass: on the document-store leg the store must have
    actually stamped the upload (``stamped_uploads``) *and* the recorded
    as-stored hash must differ from the provenance hash -- otherwise the
    equality below is satisfied by a store that happened to retain the bytes
    verbatim and proves nothing. On the filesystem leg the two hashes must be
    equal, which is the regression guard that this binding's behavior did not
    move.
    """
    _services, _config, handle = mpi_vault
    src = _write_office_file(tmp_path, "probe_alpha.docx", "Alpha")
    delivered = _sha256(src.read_bytes())

    result = _parse(await ingest_document(_VAULT_ID, str(src), "docx"))

    assert "error" not in result, result
    doc = await _services.graph_store.get_document(result["id"])
    assert doc.source_content_hash == delivered, (
        "source_content_hash must identify the bytes the caller delivered"
    )
    if handle.fake_client is not None:
        assert handle.fake_client.stamped_uploads >= 1, (
            "the document store must have rewritten the upload; without a real "
            "divergence this test asserts nothing"
        )
        assert doc.stored_content_hash != delivered, (
            "the as-stored hash must record the rewritten copy, not the delivered bytes"
        )
        assert doc.stored_content_hash == handle.fake_client.hash_source_bytes(
            _VAULT_ID, doc.source_path
        )
    else:
        assert doc.stored_content_hash == delivered, (
            "the filesystem binding stores what it was given; the two hashes must agree"
        )


@requires_docx
async def test_mpi_010_identical_reingest_dedups_without_a_second_stored_copy(mpi_vault, tmp_path):
    """MPI-010: re-ingesting byte-identical content is rejected as a duplicate
    and leaves no additional stored copy behind.

    Anti-coincidental-pass: the duplicate error alone does not prove the second
    defect is fixed -- today's retain path uploads a disambiguated second copy
    *before* duplicate detection ever runs. The upload counter is the assertion
    with teeth: it must be unchanged across the second call. It needs its
    positive control to mean anything, though -- "unchanged" is also true of an
    implementation that never retained anything at all -- so the first ingest is
    asserted to have actually stored a copy before the counter is read.
    """
    _services, _config, handle = mpi_vault
    src = _write_office_file(tmp_path, "probe_alpha.docx", "Alpha")

    first = _parse(await ingest_document(_VAULT_ID, str(src), "docx"))
    assert "error" not in first, first

    uploads_before = handle.fake_client.source_uploads if handle.fake_client else None
    stored_before = dict(handle.fake_client.sources) if handle.fake_client else None
    if handle.fake_client is not None:
        assert uploads_before >= 1, (
            "positive control: the first ingest must have stored a copy, or "
            "'the counter did not move' is satisfied by storing nothing at all"
        )
        assert first["source_path"] in stored_before

    second = _parse(await ingest_document(_VAULT_ID, str(src), "docx"))

    assert second.get("error") == "duplicate_content", second
    assert second["detail"]["existing_document_id"] == first["id"]
    assert len(await _services.graph_store.list_all_documents()) == 1

    if handle.fake_client is not None:
        assert handle.fake_client.source_uploads == uploads_before, (
            "a byte-identical re-ingest must not write another copy to the store"
        )
        assert set(handle.fake_client.sources) == set(stored_before), (
            "no additional stored path may appear on a duplicate re-ingest"
        )


@requires_docx
async def test_mpi_011_identical_bytes_under_a_new_name_still_dedup(mpi_vault, tmp_path):
    """MPI-011: identical bytes delivered under a different filename are still
    rejected as a duplicate -- duplicate detection is hash-keyed, not
    name-keyed, under either binding.

    Anti-coincidental-pass: assert the reported hash is the *delivered* one. An
    implementation that deduped on the as-stored hash could not match here at
    all (the store mints a fresh copy per upload), so a passing assertion on the
    hash value pins which hash the rejection was keyed to.
    """
    _services, _config, _handle = mpi_vault
    src = _write_office_file(tmp_path, "probe_alpha.docx", "Alpha")
    delivered = _sha256(src.read_bytes())

    first = _parse(await ingest_document(_VAULT_ID, str(src), "docx"))
    assert "error" not in first, first

    renamed = src.parent / "probe_beta.docx"
    renamed.write_bytes(src.read_bytes())
    second = _parse(await ingest_document(_VAULT_ID, str(renamed), "docx"))

    assert second.get("error") == "duplicate_content", second
    assert second["detail"]["source_content_hash"] == delivered
    assert len(await _services.graph_store.list_all_documents()) == 1


async def _audit(services, config):
    """Run the real source-file integrity audit over this vault.

    Built directly over the fixture's store pair: the profile-invariance vault is
    initialized without a registry service, so it carries no maintenance service
    of its own. Running the production audit rather than re-deriving its
    comparator here is the point -- a test that recomputed "what the audit would
    say" would keep passing through a regression in the comparator itself.
    """
    from sage.services.maintenance import MaintenanceService

    return await MaintenanceService(
        vault_id=config.vault.id,
        graph_store=services.graph_store,
        config=config,
        registry_service=None,
        content_store=StubContentStore(),
    ).verify_vault_source_files(check_hashes=True)


def _alter_retained_copy(config, handle, retained_path: str) -> bytes:
    """Alter a retained copy out of band and return what it held before.

    The condition the binding contract excludes and the source-file integrity
    audit exists to surface. Asserts the alteration took effect, so a test built
    on it cannot silently exercise the ordinary unaltered path.
    """
    if handle.fake_client is not None:
        original = handle.fake_client.sources[retained_path]
        handle.fake_client.sources[retained_path] = original + b" altered out of band"
        assert handle.fake_client.sources[retained_path] != original
    else:
        on_disk = Path(config.vault.storage_root) / retained_path
        original = on_disk.read_bytes()
        on_disk.write_bytes(original + b" altered out of band")
        assert on_disk.read_bytes() != original
    return original


@requires_docx
async def test_mpi_012_unforced_redelivery_leaves_an_altered_copy_untouched(mpi_vault, tmp_path):
    """MPI-012: an unforced re-delivery of the original bytes declines as a
    duplicate without disturbing a copy that drifted out of band, so the
    integrity audit still reports that copy as mismatched.

    Pins the deliberate trade in ``_retain_or_reuse``: on the unforced path the
    reuse reads the recorded provenance rather than re-reading the retained
    bytes, because confirming would cost a whole-file read (a network
    round-trip for a remote store) on every ingest.

    Anti-coincidental-pass: the copy is genuinely altered first (asserted inside
    the helper), so this is the drifted branch and not the ordinary
    identical-copy path; and the audit is actually run afterwards rather than
    asserted about in prose. An implementation that overwrote the drifted copy
    here would leave the audit clean and fail the final assertion -- which is
    the whole claim, since a silent repair would also erase the operator's
    evidence that something wrote to the store.
    """
    services, config, handle = mpi_vault
    src = _write_office_file(tmp_path, "probe_alpha.docx", "Alpha")
    first = _parse(await ingest_document(_VAULT_ID, str(src), "docx"))
    assert "error" not in first, first
    retained_path = first["source_path"]
    _alter_retained_copy(config, handle, retained_path)
    uploads_before = handle.fake_client.source_uploads if handle.fake_client else None

    second = _parse(await ingest_document(_VAULT_ID, str(src), "docx"))

    assert second.get("error") == "duplicate_content", second
    if handle.fake_client is not None:
        assert handle.fake_client.source_uploads == uploads_before, (
            "an unforced duplicate must not write to the store at all"
        )

    report = await _audit(services, config)
    assert report.summary["hash_mismatch"] == 1, (
        "the drifted copy must still be reported; a silent repair would erase the "
        "operator's only evidence that something else wrote to the store"
    )


@requires_docx
async def test_mpi_013_forced_redelivery_does_not_launder_a_drifted_copy(mpi_vault, tmp_path):
    """MPI-013: a forced re-delivery over a copy that drifted out of band does
    not adopt the drift as the expected state -- the audit still reports the
    copy as mismatched afterwards.

    A forced re-ingest reuses the retained copy like any other re-delivery, so
    it does not repair the drift; what it must not do is *hide* it. The
    as-stored digest is refreshed only when this call actually wrote the source,
    so a reuse leaves the recorded digest describing the copy that was last
    written and the mismatch stays visible.

    Restoring a drifted copy in place is a capability this seam does not offer:
    writing bytes back to a chosen path is not expressible through the port, and
    a forced re-ingest cannot stand in for it -- the binding sees only that the
    bytes at its target differ from the ones offered, which is indistinguishable
    from a name collision, so it would disambiguate rather than overwrite.

    Anti-coincidental-pass: the audit is run and asserted *red*, which is the
    whole claim; and the drift is asserted to have taken effect first, so the
    red verdict cannot come from an unaltered copy. The rival this excludes --
    refreshing the digest unconditionally on the force branch -- leaves the
    audit green here while changing nothing else observable, which is precisely
    what makes it invisible without this assertion.
    """
    services, config, handle = mpi_vault
    src = _write_office_file(tmp_path, "probe_alpha.docx", "Alpha")
    first = _parse(await ingest_document(_VAULT_ID, str(src), "docx"))
    assert "error" not in first, first
    retained_path = first["source_path"]
    _alter_retained_copy(config, handle, retained_path)

    second = _parse(await ingest_document(_VAULT_ID, str(src), "docx", force=True))

    assert "error" not in second, second
    assert second["source_path"] == retained_path

    report = await _audit(services, config)
    assert report.summary["hash_mismatch"] == 1, (
        "a forced re-delivery must not adopt an out-of-band change as the "
        "expected state; the drift has to stay reported"
    )
    entry = report.entries[0]
    assert entry.expected_content_hash == second["stored_content_hash"]


@requires_docx
async def test_mpi_014_redelivering_a_disambiguated_document_leaves_its_copy_intact(
    mpi_vault, tmp_path
):
    """MPI-014: re-delivering the bytes of a document that a name collision moved
    aside leaves that document's stored copy untouched.

    The path a name-keyed reuse check cannot see. Reuse that asked only "what is
    recorded at the un-disambiguated destination?" would read the *other*
    document's hash, fall through to a retain, and -- because the disambiguating
    suffix is derived from the delivered bytes -- rewrite this document's copy
    with a freshly stamped one, while duplicate detection then declined the
    ingest and left the record describing the copy that was just replaced. The
    audit would report that document mismatched forever after, on a call that
    told the caller nothing changed.

    Anti-coincidental-pass: the collision is constructed and asserted (the two
    documents really do land on different paths under the same basename), and
    the audit is run at the end. Asserting only the duplicate envelope would
    pass against the defect, since the rejection happens either way -- the
    stored-bytes comparison and the audit are what discriminate.
    """
    services, config, handle = mpi_vault
    if handle.fake_client is None:
        pytest.skip("the rewrite-on-retain condition is specific to a restamping store")

    first_src = _write_office_file(tmp_path, "x.docx", "Alpha")
    doc_a = _parse(await ingest_document(_VAULT_ID, str(first_src), "docx"))
    assert "error" not in doc_a, doc_a

    second_dir = tmp_path / "second"
    second_dir.mkdir()
    second_src = _write_office_file(second_dir, "x.docx", "Bravo with different content")
    doc_b = _parse(await ingest_document(_VAULT_ID, str(second_src), "docx"))
    assert "error" not in doc_b, doc_b
    assert doc_b["source_path"] != doc_a["source_path"], (
        "the collision must actually disambiguate, or this exercises the base path"
    )

    stored_before = handle.fake_client.sources[doc_b["source_path"]]
    uploads_before = handle.fake_client.source_uploads

    redelivered = _parse(await ingest_document(_VAULT_ID, str(second_src), "docx"))

    assert redelivered.get("error") == "duplicate_content", redelivered
    assert handle.fake_client.sources[doc_b["source_path"]] == stored_before, (
        "re-delivering unchanged bytes must not rewrite the disambiguated copy"
    )
    assert handle.fake_client.source_uploads == uploads_before

    report = await _audit(services, config)
    assert report.summary["hash_mismatch"] == 0, report.summary


@requires_docx
async def test_mpi_015_resident_source_reingest_dedups_rather_than_duplicating(mpi_vault, tmp_path):
    """MPI-015: re-ingesting a document by its vault-relative path, when the
    bytes live only on the store, is rejected as a duplicate rather than landing
    a second document on the same retained file.

    The post-restart cloud condition that branch exists for. Nothing is
    delivered on this call, so the provenance of the path is inherited from the
    record that established it; deriving it from the stored copy instead would
    give the re-projection a different identity from the document it is
    re-projecting -- and with no unique index on ``source_path``, duplicate
    detection missing means an insert, not an error.

    Anti-coincidental-pass: asserts the document count, not just the envelope. A
    document count of one is what separates "rejected as a duplicate" from
    "silently inserted alongside"; the error code alone cannot, because the
    defective path raises no error at all.
    """
    _services, _config, handle = mpi_vault
    src = _write_office_file(tmp_path, "resident.docx", "Alpha")
    first = _parse(await ingest_document(_VAULT_ID, str(src), "docx"))
    assert "error" not in first, first
    rel = first["source_path"]
    if handle.fake_client is not None:
        assert rel in handle.fake_client.sources, "the bytes must be on the store"

    second = _parse(await ingest_document(_VAULT_ID, rel, "docx"))

    assert second.get("error") == "duplicate_content", second
    docs = await _services.graph_store.list_all_documents()
    assert len(docs) == 1, (
        f"a resident-path re-ingest must not insert a second document; got "
        f"{[(d.id, d.source_path) for d in docs]}"
    )


@requires_docx
async def test_mpi_016_reuse_falls_through_when_the_stored_copy_is_gone(mpi_vault, tmp_path):
    """MPI-016: when the recorded copy has vanished from the store, a
    re-delivery retains the bytes again instead of reusing a path that resolves
    to nothing.

    The third fall-through condition in ``_retain_or_reuse`` -- the other two
    (no document with these bytes, and a forced ingest) are covered by MPI-009
    and MPI-013. Without the presence check a document would be "reused" onto a
    path holding no bytes, which every later read and the integrity audit would
    then report as missing.

    Anti-coincidental-pass: the copy is deleted and its absence asserted, so the
    guard is the only thing standing between this call and a dangling reuse.
    Mutating the guard away leaves the store with no copy at that path and fails
    the final presence assertion.
    """
    _services, config, handle = mpi_vault
    src = _write_office_file(tmp_path, "vanishing.docx", "Alpha")
    first = _parse(await ingest_document(_VAULT_ID, str(src), "docx"))
    assert "error" not in first, first
    retained_path = first["source_path"]

    if handle.fake_client is not None:
        del handle.fake_client.sources[retained_path]
        assert retained_path not in handle.fake_client.sources
    else:
        on_disk = Path(config.vault.storage_root) / retained_path
        on_disk.unlink()
        assert not on_disk.exists()

    second = _parse(await ingest_document(_VAULT_ID, str(src), "docx", force=True))

    assert "error" not in second, second
    if handle.fake_client is not None:
        assert second["source_path"] in handle.fake_client.sources, (
            "the vanished copy must be retained again, not reused as a dangling path"
        )
    else:
        assert (Path(config.vault.storage_root) / second["source_path"]).exists()


def test_mpi_007_path_parameter_docstrings_state_server_local_contract():
    """MPI-007: every path-bearing tool documents that its path parameter
    resolves on the machine running the SAGE server process."""
    contract = "machine running the SAGE server process"
    for tool in (
        ingest_document,
        bulk_ingest_document,
        get_document,
        read_projection,
        list_directory,
    ):
        # Collapse whitespace so the check is about content, not where the
        # docstring happens to wrap.
        doc = " ".join((inspect.getdoc(tool) or "").split())
        assert contract in doc, f"{tool.__name__} docstring lacks the server-local path contract"
