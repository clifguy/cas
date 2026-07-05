"""Docx source adapter: parses Word document styles, numbering, and tables.

Extracts heading hierarchy from paragraph styles, computes heading number
prefixes from Word numbering definitions, renders tables as pipe-delimited
text, and resolves cross-reference fields via cached display values.

Handles both Word documents (.docx) and Word templates (.dotx). The two
formats share WordprocessingML body structure; they differ only in the
main part's OPC content type. python-docx rejects.dotx at load time, so
the adapter rewrites the content-type entry in a temp copy before loading.

For.dotx files the adapter additionally surfaces a structured style
inventory in metadata["template_style_inventory"] and emits namespaced
tags via metadata["adapter_tags"] so agentic queries can discover
templates by defined style and by auto-numbering.

Computes SHA-256 of raw file bytes for content_hash.
"""

import hashlib
import shutil
import tempfile
import zipfile
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from lxml import etree

from sage.source_adapters.base import (
    HeadingNode,
    ProjectionResult,
    SourceAdapter,
    extract_adr_id_from_filename,
)

# OPC content types for the main document part.
_DOCX_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"
)
_DOTX_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.template.main+xml"
)

# Tag namespace prefixes owned by this adapter (for force-reingest stripping
# in the ingestion service). All begin with "template:" so non-template
# .docx files never produce tags in these namespaces.
_ADAPTER_TAG_PREFIXES = ["template:"]

# Word namespace
_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

# Default Word heading styles mapped to heading levels
_DEFAULT_STYLE_MAP: dict[str, int] = {f"Heading {i}": i for i in range(1, 10)}


class _NumberingEngine:
    """Computes rendered heading number prefixes from Word numbering definitions.

    Parses w:abstractNum definitions from the document's numbering part,
    tracks per-level counters, and formats composite prefixes using the
    lvlText template (e.g., "%1.%2" -> "1.2").
    """

    def __init__(self, numbering_part) -> None:
        # abstractNumId -> {ilvl -> {start, numFmt, lvlText}}
        self._abstract_defs: dict[int, dict[int, dict]] = {}
        # numId -> abstractNumId
        self._num_map: dict[int, int] = {}
        # (numId, ilvl) -> current counter value
        self._counters: dict[tuple[int, int], int] = {}

        if numbering_part is None:
            return

        numbering_elm = numbering_part.numbering_definitions._numbering

        for abstract in numbering_elm.findall(qn("w:abstractNum")):
            abs_id = int(abstract.get(qn("w:abstractNumId")))
            levels: dict[int, dict] = {}
            for lvl in abstract.findall(qn("w:lvl")):
                ilvl = int(lvl.get(qn("w:ilvl")))
                start_elem = lvl.find(qn("w:start"))
                fmt_elem = lvl.find(qn("w:numFmt"))
                text_elem = lvl.find(qn("w:lvlText"))
                levels[ilvl] = {
                    "start": int(start_elem.get(qn("w:val"))) if start_elem is not None else 1,
                    "numFmt": fmt_elem.get(qn("w:val")) if fmt_elem is not None else "decimal",
                    "lvlText": text_elem.get(qn("w:val")) if text_elem is not None else "",
                }
            self._abstract_defs[abs_id] = levels

        for num in numbering_elm.findall(qn("w:num")):
            num_id = int(num.get(qn("w:numId")))
            abs_ref = num.find(qn("w:abstractNumId"))
            if abs_ref is not None:
                self._num_map[num_id] = int(abs_ref.get(qn("w:val")))

    def next_number(self, num_id: int, ilvl: int) -> str:
        """Advance counter for (numId, ilvl) and return formatted prefix."""
        abs_id = self._num_map.get(num_id)
        if abs_id is None:
            return ""
        levels = self._abstract_defs.get(abs_id)
        if levels is None:
            return ""
        level_def = levels.get(ilvl)
        if level_def is None:
            return ""

        # Initialize counter if needed
        key = (num_id, ilvl)
        if key not in self._counters:
            self._counters[key] = level_def["start"]
        else:
            self._counters[key] += 1

        # Reset all child levels
        for child_ilvl in levels:
            if child_ilvl > ilvl:
                child_key = (num_id, child_ilvl)
                if child_key in self._counters:
                    del self._counters[child_key]

        # Build formatted text from lvlText template
        lvl_text = level_def["lvlText"]
        result = lvl_text
        for ref_ilvl in range(ilvl + 1):
            placeholder = f"%{ref_ilvl + 1}"
            if placeholder in result:
                ref_key = (num_id, ref_ilvl)
                counter_val = self._counters.get(ref_key, levels.get(ref_ilvl, {}).get("start", 1))
                ref_fmt = levels.get(ref_ilvl, {}).get("numFmt", "decimal")
                result = result.replace(placeholder, _format_number(counter_val, ref_fmt))

        return result


def _format_style_name_with_numbering(entry: dict) -> str:
    """Render a style inventory entry for the synthesized template text.

    Appends " (auto-numbered)" when the entry carries active numbering
    so the projected text preserves this discriminating fact for both
    retrieval (BM25/vector will match on "auto-numbered") and human
    review of the text blob.
    """
    name = entry["name"]
    if entry.get("has_numbering"):
        return f"{name} (auto-numbered)"
    return name


def _format_number(value: int, num_fmt: str) -> str:
    """Format a counter value according to Word's numFmt."""
    if num_fmt == "decimal":
        return str(value)
    elif num_fmt == "upperRoman":
        return _to_roman(value)
    elif num_fmt == "lowerRoman":
        return _to_roman(value).lower()
    elif num_fmt == "upperLetter":
        return _to_letter(value).upper()
    elif num_fmt == "lowerLetter":
        return _to_letter(value)
    else:
        return str(value)


def _to_roman(n: int) -> str:
    """Convert integer to Roman numeral string."""
    vals = [
        (1000, "M"),
        (900, "CM"),
        (500, "D"),
        (400, "CD"),
        (100, "C"),
        (90, "XC"),
        (50, "L"),
        (40, "XL"),
        (10, "X"),
        (9, "IX"),
        (5, "V"),
        (4, "IV"),
        (1, "I"),
    ]
    result = ""
    for val, numeral in vals:
        while n >= val:
            result += numeral
            n -= val
    return result


def _to_letter(n: int) -> str:
    """Convert integer to lowercase letter (1=a, 2=b,..., 26=z, 27=aa)."""
    result = ""
    while n > 0:
        n -= 1
        result = chr(ord("a") + (n % 26)) + result
        n //= 26
    return result


# Common English stop words for title key-term extraction.
_STOP_WORDS = frozenset(
    "a an the and or but in on of to for with by from at is are was were "
    "be been being have has had do does did will would shall should may might "
    "can could this that these those it its he she they we you i my our his "
    "her their your not no nor so if as than too very also about above after "
    "before between into through during each all any both few more most other "
    "some such only same then there here when where how what which who whom "
    "why".split()
)

_MAX_KEY_TERMS = 6


def _extract_key_terms(text: str) -> str:
    """Extract key terms from text by filtering stop words.

    Returns up to _MAX_KEY_TERMS content words joined by spaces.
    Used as a last-resort title when no Title style or filename is available.
    """
    import re

    words = re.findall(r"[A-Za-z0-9]+(?:'[a-z]+)?", text)
    terms = [w for w in words if w.lower() not in _STOP_WORDS and len(w) > 1]
    return " ".join(terms[:_MAX_KEY_TERMS]) if terms else "Untitled"


def _extract_paragraph_text(p_elem) -> str:
    """Extract visible text from a paragraph element.

    Walks w:r/w:t elements and collects text. Skips w:instrText elements
    so that cross-reference field instructions are excluded while their
    cached display values (in w:t runs between separate/end markers) are
    included.

    Also handles w:fldSimple elements whose child runs contain display text.
    """
    parts: list[str] = []
    for child in p_elem:
        tag = etree.QName(child.tag).localname if isinstance(child.tag, str) else ""
        if tag == "r":
            # Regular run: collect w:t, skip w:instrText
            for run_child in child:
                rc_tag = (
                    etree.QName(run_child.tag).localname if isinstance(run_child.tag, str) else ""
                )
                if rc_tag == "t" and run_child.text:
                    parts.append(run_child.text)
        elif tag == "fldSimple":
            # Simple field: collect text from child runs
            for r in child.findall(qn("w:r")):
                for t in r.findall(qn("w:t")):
                    if t.text:
                        parts.append(t.text)
    return "".join(parts)


class DocxAdapter(SourceAdapter):
    """Source adapter for Microsoft Word (.docx) documents and (.dotx) templates."""

    # 0.3.0: chunks indexed with heading-context (heading_path embedded with
    # content, plus FTS index on heading_path) — see ingestion._stage2_indexing.
    # 0.4.0: chunker emits one chunk per heading regardless of body content,
    # so heading paragraphs whose next paragraph is another heading still
    # have their heading_path indexed (Word-Find equivalence).
    VERSION = "0.4.0"
    EXTENSIONS = [".docx", ".dotx"]

    async def project(self, source_path: Path, config: dict | None = None) -> ProjectionResult:
        raw_bytes = source_path.read_bytes()
        content_hash = hashlib.sha256(raw_bytes).hexdigest()

        is_template = source_path.suffix.lower() == ".dotx"
        doc = self._open_document(source_path, is_template)
        style_map = self._build_style_map(config)
        style_id_to_name = self._build_style_id_to_name(doc)

        # Initialize numbering engine
        numbering_part = None
        try:
            numbering_part = doc.part.numbering_part
        except Exception:  # noqa: S110 -- optional docx feature; absence is expected
            pass
        engine = _NumberingEngine(numbering_part)

        headings: list[HeadingNode] = []
        text_parts: list[str] = []
        stack: list[tuple[int, str]] = []
        current_content_lines: list[str] = []
        current_heading_idx = -1
        title_style_text: str | None = None
        first_body_para: str | None = None

        # Walk document body children in order (paragraphs and tables)
        for element in doc.element.body:
            tag = etree.QName(element.tag).localname if isinstance(element.tag, str) else ""

            if tag == "p":
                para_text = _extract_paragraph_text(element)

                # Determine style name
                style_name = self._get_style_name(element, style_id_to_name)
                level = style_map.get(style_name) if style_name else None

                # Capture first paragraph with Title style
                if title_style_text is None and style_name == "Title" and para_text.strip():
                    title_style_text = para_text.strip()

                if level is not None:
                    # Flush content to previous heading
                    if current_heading_idx >= 0:
                        headings[current_heading_idx].content = "\n".join(
                            current_content_lines
                        ).strip()
                    current_content_lines = []

                    # Check for numbering
                    prefix = self._get_numbering_prefix(element, engine)
                    heading_text = f"{prefix} {para_text}" if prefix else para_text

                    # Update heading stack
                    while stack and stack[-1][0] >= level:
                        stack.pop()
                    stack.append((level, heading_text))

                    path = " > ".join(h[1] for h in stack)

                    headings.append(
                        HeadingNode(
                            level=level,
                            text=heading_text,
                            path=path,
                            content="",
                        )
                    )
                    current_heading_idx = len(headings) - 1
                    text_parts.append(heading_text)
                else:
                    if para_text.strip():
                        current_content_lines.append(para_text)
                        text_parts.append(para_text)
                        # Capture first non-heading, non-title body paragraph
                        if first_body_para is None and style_name != "Title":
                            first_body_para = para_text.strip()

            elif tag == "tbl":
                table_text = self._render_table(element)
                if table_text:
                    current_content_lines.append(table_text)
                    text_parts.append(table_text)

        # Flush final heading content
        if current_heading_idx >= 0:
            headings[current_heading_idx].content = "\n".join(current_content_lines).strip()

        title = self._extract_title(
            title_style_text,
            first_body_para,
            source_path,
        )

        source_mtime = datetime.fromtimestamp(source_path.stat().st_mtime, tz=timezone.utc)

        metadata: dict = {"source_modified_at": source_mtime.isoformat()}
        adapter_tier3 = extract_adr_id_from_filename(source_path.stem)
        if adapter_tier3 is not None:
            metadata["adapter_tier3_metadata"] = adapter_tier3

        projected_text = "\n".join(text_parts)

        # Template-only enrichment: style inventory, namespaced tags, and
        # a synthesized prose projection derived from the style surface.
        # Gated on extension, not on content heuristics, so document
        # metadata and text stay lean for.docx files.
        if is_template:
            inventory = self._build_template_style_inventory(doc)
            metadata["template_style_inventory"] = inventory
            metadata["adapter_tags"] = self._build_template_tags(inventory)
            metadata["adapter_tag_prefixes"] = list(_ADAPTER_TAG_PREFIXES)

            # Prepend a synthesized prose description of the template's
            # style surface. Templates are style-rich and content-thin; a
            # raw body parse usually produces near-zero text, which
            # blocks retrieval indexing. Prepending (rather than
            # appending) puts the template-identity signal at the head
            # of the text so the embedding captures "Word template with
            # these styles" as the dominant topic, and so BM25 matches
            # on custom style names surface with strong weights.
            style_surface_text = self._synthesize_template_text(title, inventory)
            if projected_text.strip():
                projected_text = style_surface_text + "\n\n" + projected_text
            else:
                projected_text = style_surface_text

        return ProjectionResult(
            text=projected_text,
            headings=headings,
            content_hash=content_hash,
            adapter_version=self.VERSION,
            title=title,
            metadata=metadata,
        )

    @staticmethod
    def _synthesize_template_text(title: str, inventory: list[dict]) -> str:
        """Build a prose description of a template's style surface.

        Designed for retrieval indexing: lists every custom style by its
        human name, notes which built-in styles the template wires to
        active auto-numbering, and identifies the artifact as a Word
        template. Built-in styles whose behavior is unchanged are
        omitted to prevent every template's text from collapsing into
        an identical list of stock Word styles.
        """
        custom = [e for e in inventory if e.get("is_custom")]
        builtin_with_numbering = [
            e for e in inventory if not e.get("is_custom") and e.get("has_numbering")
        ]

        lines: list[str] = [
            f"Microsoft Word template: {title}.",
        ]

        if custom:
            by_type: dict[str, list[dict]] = {}
            for e in custom:
                by_type.setdefault(e["type"], []).append(e)

            # Stable type ordering for deterministic projection output
            type_order = ["paragraph", "character", "table", "numbering"]
            custom_count = len(custom)
            lines.append("")
            lines.append(
                f"This template defines {custom_count} user-authored "
                f"style{'s' if custom_count != 1 else ''}:"
            )
            for t in type_order:
                entries = by_type.get(t) or []
                if not entries:
                    continue
                names = ", ".join(_format_style_name_with_numbering(e) for e in entries)
                plural = "styles" if len(entries) != 1 else "style"
                lines.append(f"  {t.capitalize()} {plural}: {names}.")

        if builtin_with_numbering:
            names = ", ".join(e["name"] for e in builtin_with_numbering)
            lines.append("")
            lines.append(f"Built-in styles carrying template-local auto-numbering: {names}.")

        return "\n".join(lines)

    def _open_document(self, source_path: Path, is_template: bool) -> Document:
        """Load a.docx or.dotx into a python-docx Document.

        python-docx rejects.dotx at load time because the main part's OPC
        content type is the template flavor, not the document flavor. We
        work around this by copying the ZIP into a temp file and rewriting
        the [Content_Types].xml entry. The body XML and all other parts are
        unchanged, so the rest of the adapter is unaffected.
        """
        if not is_template:
            return Document(str(source_path))

        tmp_dir = Path(tempfile.mkdtemp(prefix="sage_dotx_"))
        try:
            shadow = tmp_dir / "shadow.docx"
            with zipfile.ZipFile(source_path, "r") as z_in:
                with zipfile.ZipFile(shadow, "w", zipfile.ZIP_DEFLATED) as z_out:
                    for item in z_in.namelist():
                        data = z_in.read(item)
                        if item == "[Content_Types].xml":
                            data = data.replace(
                                _DOTX_CONTENT_TYPE.encode("utf-8"),
                                _DOCX_CONTENT_TYPE.encode("utf-8"),
                            )
                        z_out.writestr(item, data)
            return Document(str(shadow))
        finally:
            # python-docx has loaded the file into memory by the time
            # Document() returns, so the temp dir is safe to remove.
            shutil.rmtree(tmp_dir, ignore_errors=True)

    def _build_template_style_inventory(self, doc: Document) -> list[dict]:
        """Build a structured inventory of the template's style surface.

        Returns a list of entries, each with keys:
        - id: XML style ID (e.g., "Heading1")
        - name: human-readable style name (e.g., "Heading 1")
        - type: one of "paragraph", "character", "table", "numbering"
        - based_on: XML style ID this style inherits from, or None
        - has_numbering: True iff a paragraph style carries a numPr whose
          numId resolves to an active abstract numbering definition
        - is_custom: True iff the style element carries w:customStyle="1"
        - numbering_detail: resolved numbering definition (dict) when
          has_numbering is True, otherwise None. See _resolve_numbering.
        """
        # Resolve numbering references to know which numIds are "active"
        # and to have the full numbering map available for numbering_detail.
        num_map, abstract_defs, num_overrides = self._build_numbering_maps(doc)
        active_num_ids = set(num_map.keys())

        type_map = {
            "PARAGRAPH": "paragraph",
            "CHARACTER": "character",
            "TABLE": "table",
            "LIST": "numbering",
        }

        inventory: list[dict] = []
        for style in doc.styles:
            if not style.style_id or not style.name:
                continue
            type_name = getattr(style.type, "name", "") or ""
            style_type = type_map.get(type_name, "paragraph")

            # basedOn (inheritance parent) -- read from XML directly; python-docx
            # exposes this inconsistently across style types.
            based_on: str | None = None
            based_on_elem = style.element.find(qn("w:basedOn"))
            if based_on_elem is not None:
                based_on = based_on_elem.get(qn("w:val"))

            has_numbering = False
            numbering_detail: dict | None = None
            style_num_ref = self._extract_style_numbering_reference(style)
            if style_type == "paragraph" and style_num_ref is not None:
                num_id_val, ilvl_val = style_num_ref
                if num_id_val != 0 and num_id_val in active_num_ids:
                    has_numbering = True
                    numbering_detail = self._resolve_numbering(
                        num_id_val,
                        ilvl_val,
                        num_map,
                        abstract_defs,
                        num_overrides,
                    )

            is_custom = style.element.get(qn("w:customStyle")) == "1"

            inventory.append(
                {
                    "id": style.style_id,
                    "name": style.name,
                    "type": style_type,
                    "based_on": based_on,
                    "has_numbering": has_numbering,
                    "is_custom": is_custom,
                    "numbering_detail": numbering_detail,
                }
            )

        return inventory

    @staticmethod
    def _build_numbering_maps(
        doc: Document,
    ) -> tuple[dict[int, int], dict[int, dict[int, dict]], dict[int, dict[int, dict]]]:
        """Parse the document's numbering part into three lookup maps.

        Returns:
        - num_map: numId -> abstractNumId
        - abstract_defs: abstractNumId -> {ilvl -> {num_fmt, lvl_text}}
        - num_overrides: numId -> {ilvl -> {num_fmt, lvl_text}} (from
          <w:lvlOverride><w:lvl>...</w:lvl></w:lvlOverride> entries).
          Overrides may be partial -- only the fields set in the XML
          are present.
        """
        num_map: dict[int, int] = {}
        abstract_defs: dict[int, dict[int, dict]] = {}
        num_overrides: dict[int, dict[int, dict]] = {}

        try:
            numbering_part = doc.part.numbering_part
        except Exception:
            numbering_part = None
        if numbering_part is None:
            return num_map, abstract_defs, num_overrides

        numbering_elm = numbering_part.numbering_definitions._numbering

        for abstract in numbering_elm.findall(qn("w:abstractNum")):
            abs_id_attr = abstract.get(qn("w:abstractNumId"))
            if abs_id_attr is None:
                continue
            try:
                abs_id = int(abs_id_attr)
            except ValueError:
                continue
            levels: dict[int, dict] = {}
            for lvl in abstract.findall(qn("w:lvl")):
                ilvl_attr = lvl.get(qn("w:ilvl"))
                if ilvl_attr is None:
                    continue
                try:
                    ilvl = int(ilvl_attr)
                except ValueError:
                    continue
                fmt_elem = lvl.find(qn("w:numFmt"))
                text_elem = lvl.find(qn("w:lvlText"))
                levels[ilvl] = {
                    "num_fmt": (fmt_elem.get(qn("w:val")) if fmt_elem is not None else "decimal"),
                    "lvl_text": (text_elem.get(qn("w:val")) if text_elem is not None else ""),
                }
            abstract_defs[abs_id] = levels

        for num in numbering_elm.findall(qn("w:num")):
            num_id_attr = num.get(qn("w:numId"))
            if num_id_attr is None:
                continue
            try:
                num_id = int(num_id_attr)
            except ValueError:
                continue
            abs_ref = num.find(qn("w:abstractNumId"))
            if abs_ref is not None:
                try:
                    num_map[num_id] = int(abs_ref.get(qn("w:val")))
                except (TypeError, ValueError):
                    continue

            # Capture lvlOverride entries. A lvlOverride may contain a
            # <w:lvl> that replaces the abstract's level definition for
            # this num only (e.g., to force numFmt=none for suppression).
            overrides: dict[int, dict] = {}
            for ovr in num.findall(qn("w:lvlOverride")):
                ovr_ilvl_attr = ovr.get(qn("w:ilvl"))
                if ovr_ilvl_attr is None:
                    continue
                try:
                    ovr_ilvl = int(ovr_ilvl_attr)
                except ValueError:
                    continue
                lvl = ovr.find(qn("w:lvl"))
                if lvl is None:
                    continue
                level_patch: dict = {}
                fmt_elem = lvl.find(qn("w:numFmt"))
                if fmt_elem is not None:
                    level_patch["num_fmt"] = fmt_elem.get(qn("w:val"))
                text_elem = lvl.find(qn("w:lvlText"))
                if text_elem is not None:
                    level_patch["lvl_text"] = text_elem.get(qn("w:val"))
                if level_patch:
                    overrides[ovr_ilvl] = level_patch
            if overrides:
                num_overrides[num_id] = overrides

        return num_map, abstract_defs, num_overrides

    @staticmethod
    def _extract_style_numbering_reference(style) -> tuple[int, int] | None:
        """Read a paragraph style's numPr (numId, ilvl).

        Returns None if the style has no numPr. Absent ilvl defaults
        to 0 per OOXML convention. A numId of 0 is returned as (0, 0)
        and the caller is expected to treat it as "numbering suppressed
        on this style".
        """
        pPr = style.element.find(qn("w:pPr"))
        if pPr is None:
            return None
        numPr = pPr.find(qn("w:numPr"))
        if numPr is None:
            return None
        num_id_elem = numPr.find(qn("w:numId"))
        if num_id_elem is None:
            return None
        try:
            num_id_val = int(num_id_elem.get(qn("w:val")))
        except (TypeError, ValueError):
            return None
        ilvl_elem = numPr.find(qn("w:ilvl"))
        ilvl_val = 0
        if ilvl_elem is not None:
            try:
                ilvl_val = int(ilvl_elem.get(qn("w:val")))
            except (TypeError, ValueError):
                ilvl_val = 0
        return num_id_val, ilvl_val

    @staticmethod
    def _resolve_numbering(
        num_id: int,
        ilvl: int,
        num_map: dict[int, int],
        abstract_defs: dict[int, dict[int, dict]],
        num_overrides: dict[int, dict[int, dict]],
    ) -> dict:
        """Resolve a (num_id, ilvl) reference into a concrete definition.

        Looks up the abstract's level definition, applies any
        lvlOverride patches from the num, and returns the flattened
        numbering_detail dict. `suppressed` is True iff the resolved
        num_fmt is "none".
        """
        abstract_num_id = num_map.get(num_id)
        base_def: dict = {}
        if abstract_num_id is not None:
            levels = abstract_defs.get(abstract_num_id) or {}
            base_def = dict(levels.get(ilvl) or {})

        # Apply any override at this ilvl
        override = (num_overrides.get(num_id) or {}).get(ilvl)
        if override:
            base_def.update(override)

        num_fmt = base_def.get("num_fmt") or "decimal"
        lvl_text = base_def.get("lvl_text") or ""
        return {
            "num_id": num_id,
            "abstract_num_id": abstract_num_id if abstract_num_id is not None else -1,
            "ilvl": ilvl,
            "num_fmt": num_fmt,
            "lvl_text": lvl_text,
            "suppressed": num_fmt == "none",
        }

    @staticmethod
    def _build_template_tags(inventory: list[dict]) -> list[str]:
        """Build the namespaced tag list for a template's style inventory.

        `template:style:*` is restricted to custom (user-authored) styles
        so stock Word styles do not flood every template's tag set.
        `template:has_numbering:*` applies to any style carrying active
        numbering, including built-in heading styles that the template
        has wired up with its own numbering definition -- that wiring is
        a meaningful discriminator and would be lost if limited to
        custom-only.
        """
        tags: list[str] = []
        for entry in inventory:
            if entry.get("is_custom"):
                tags.append(f"template:style:{entry['name']}")
        for entry in inventory:
            if entry.get("has_numbering"):
                tags.append(f"template:has_numbering:{entry['name']}")
        return tags

    def _build_style_map(self, config: dict | None) -> dict[str, int]:
        """Build heading style map from config, merging with defaults."""
        style_map = dict(_DEFAULT_STYLE_MAP)
        if config and "heading_style_map" in config:
            style_map.update(config["heading_style_map"])
        return style_map

    def _build_style_id_to_name(self, doc: Document) -> dict[str, str]:
        """Build a mapping from style_id (XML) to style name (human)."""
        mapping: dict[str, str] = {}
        for style in doc.styles:
            if style.style_id and style.name:
                mapping[style.style_id] = style.name
        return mapping

    def _get_style_name(self, p_elem, style_id_to_name: dict[str, str]) -> str | None:
        """Get the style name for a paragraph element."""
        pPr = p_elem.find(qn("w:pPr"))
        if pPr is None:
            return None
        pStyle = pPr.find(qn("w:pStyle"))
        if pStyle is None:
            return None
        style_id = pStyle.get(qn("w:val"))
        if style_id is None:
            return None
        return style_id_to_name.get(style_id, style_id)

    def _get_numbering_prefix(self, p_elem, engine: _NumberingEngine) -> str:
        """Extract numbering prefix for a paragraph, if any."""
        pPr = p_elem.find(qn("w:pPr"))
        if pPr is None:
            return ""
        numPr = pPr.find(qn("w:numPr"))
        if numPr is None:
            return ""
        ilvl_elem = numPr.find(qn("w:ilvl"))
        numId_elem = numPr.find(qn("w:numId"))
        if ilvl_elem is None or numId_elem is None:
            return ""
        ilvl = int(ilvl_elem.get(qn("w:val")))
        num_id = int(numId_elem.get(qn("w:val")))
        if num_id == 0:
            return ""
        return engine.next_number(num_id, ilvl)

    def _render_table(self, tbl_elem) -> str:
        """Render a table element as pipe-delimited text rows."""
        rows: list[str] = []
        for tr in tbl_elem.findall(qn("w:tr")):
            cells: list[str] = []
            for tc in tr.findall(qn("w:tc")):
                # Collect text from all paragraphs in the cell
                cell_parts: list[str] = []
                for p in tc.findall(qn("w:p")):
                    cell_parts.append(_extract_paragraph_text(p))
                cells.append(" ".join(cell_parts).strip())
            rows.append("| " + " | ".join(cells) + " |")
        return "\n".join(rows)

    def _extract_title(
        self,
        title_style_text: str | None,
        first_body_para: str | None,
        source_path: Path,
    ) -> str:
        """Extract document title with priority chain.

        1. First paragraph with Word "Title" style.
        2. Filename stem (often contains a short title or document code).
        3. Key terms from the first body paragraph (stop-word filtered).
        """
        if title_style_text:
            return title_style_text
        stem = source_path.stem
        # Reject stems that are empty or just a dotfile extension
        if stem and not stem.startswith("."):
            return stem
        if first_body_para:
            return _extract_key_terms(first_body_para)
        return "Untitled"
